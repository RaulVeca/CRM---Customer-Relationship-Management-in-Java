from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\TrainingIT_site\output\documents\Ghid_recapitulare_TrainingIT.docx")

# Preset: standard_business_brief, with a restrained academic blue/teal override.
NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "1F6E78"
INK = "222222"
MUTED = "606A75"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F4F3"
PALE_GRAY = "F3F5F7"
WHITE = "FFFFFF"
GOLD = "9C6B00"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_numbering_definition(doc, num_id, abstract_id, fmt="bullet"):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, nid])


def add_body(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_bullets(doc, items, num_id=1):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        apply_numbering(p, num_id)
        r = p.add_run(item)
        set_run_font(r)


def add_steps(doc, items, start_num_id):
    add_numbering_definition(doc, start_num_id, start_num_id + 100, fmt="decimal")
    for item in items:
        p = doc.add_paragraph(style="Normal")
        apply_numbering(p, start_num_id)
        r = p.add_run(item)
        set_run_font(r)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    prevent_row_split(table.rows[0])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_table(doc, headers, rows, widths, header_fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    repeat_header(hdr)
    prevent_row_split(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=10.2, bold=(idx == 0))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = rgb(NAVY)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(TEAL)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("TRAININGIT  |  GHID DE RECAPITULARE")
    set_run_font(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run("Pagina ")
    set_run_font(r1, size=8.5, color=MUTED)
    r2 = p.add_run()
    set_run_font(r2, size=8.5, color=MUTED)
    add_field(r2, "PAGE")


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("GHID DE RECAPITULARE")
    set_run_font(r, size=11, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("TrainingIT")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Conspect explicat pentru lucrarea de licență")
    set_run_font(r, size=15, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Aplicație web full-stack pentru comercializarea cursurilor IT și administrarea integrată a relațiilor cu clienții")
    set_run_font(r, size=11.5, italic=True, color=MUTED)

    add_callout(
        doc,
        "Ideea centrală",
        "TrainingIT reunește magazinul online de cursuri și sistemul CRM într-o singură aplicație, cu același backend și același model de date. O acțiune a clientului este înregistrată o singură dată și poate declanșa automat mai multe reacții administrative.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_heading(doc, "Cum se folosește acest ghid", 2)
    add_bullets(doc, [
        "Pentru recapitulare rapidă: citește casetele „Ideea centrală”, „De reținut” și „Clarificare”.",
        "Pentru susținere: folosește secțiunile „Cum explic oral” și răspunsurile scurte din anexă.",
        "Pentru partea tehnică: urmărește legătura dintre straturi, evenimente, observatori și baza de date.",
    ])

    add_heading(doc, "Date-cheie ale soluției", 2)
    add_table(doc, ["Componentă", "Tehnologie / rol"], [
        ("Client", "Next.js 16, React 19, TypeScript 5 și Tailwind CSS 4"),
        ("Server", "Java 17 și Spring Boot 3.5"),
        ("Date", "MariaDB, JDBC și HikariCP"),
        ("Arhitectură", "REST, SSE, evenimente și șabloane de proiectare orientate pe obiecte"),
        ("AI opțional", "Anthropic Claude pentru asistent, recomandări și traducere"),
    ], [2500, 6860], header_fill=PALE_TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Document de studiu • 2026")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def add_reading_map(doc):
    add_heading(doc, "Hartă de lectură", 1)
    add_body(doc, "Conspectul urmărește capitolele discutate și păstrează numerele din lucrarea de licență. Capitolele 3 și 4 nu au fost incluse deoarece nu au făcut parte din fragmentele conspectate în conversație.")
    rows = [
        ("Cap. 1", "Context, soluție și obiective", "De ce este necesar produsul"),
        ("Cap. 2", "Problema și abordările similare", "De ce integrarea magazin–CRM este dificilă"),
        ("Cap. 5", "Arhitectura", "Cum comunică front-endul, backendul, domeniul și baza de date"),
        ("Cap. 6", "Funcționalitățile", "Ce poate face utilizatorul și ce poate face administratorul"),
        ("Cap. 7", "Tehnologiile", "Cu ce este construită aplicația și de ce"),
        ("Cap. 8", "Implementarea", "Cum sunt folosite evenimentele, observatorii și modelul de date"),
        ("Cap. 9", "Testarea", "Ce s-a verificat și cum se comportă sistemul la erori"),
        ("Cap. 10", "Scenarii", "Cum se vede soluția în utilizare reală"),
        ("Cap. 11", "Concluzii", "Rezultate și dezvoltări viitoare"),
        ("Anexe", "Glosar și răspunsuri scurte", "Pregătire pentru întrebările de la susținere"),
    ]
    add_table(doc, ["Secțiune", "Subiect", "Întrebarea la care răspunde"], rows, [1250, 3300, 4810])
    add_callout(doc, "Fir logic", "Nevoie comercială → aplicație unificată → arhitectură pe straturi → reacții automate → verificare prin scenarii și teste.")


def chapter_1(doc):
    add_heading(doc, "Capitolul 1. Introducere și obiective", 1)
    add_heading(doc, "Contextul domeniului", 2)
    add_body(doc, "Sectorul tehnologiei informației este unul dintre cele mai dinamice domenii ale pieței muncii. În acest context, învățarea continuă nu mai este doar o opțiune, ci o necesitate profesională. Cererea pentru cursuri online de programare a crescut, iar furnizorii de formare își comercializează serviciile tot mai frecvent direct pe internet.")
    add_body(doc, "O companie care vinde cursuri online trebuie să rezolve simultan două categorii de nevoi. Prima este experiența clientului: prezentarea cursurilor, înscrierea, plata și suportul. A doua este activitatea internă: cunoașterea clienților, urmărirea etapelor de vânzare, facturarea, analiza veniturilor și soluționarea problemelor.")
    add_table(doc, ["Nevoie", "Soluție obișnuită", "Exemple de informații"], [
        ("Vânzarea cursurilor", "Platformă e-commerce", "catalog, preț, recenzii, înscriere, plată"),
        ("Relația cu clientul", "Sistem CRM", "contacte, leaduri, achiziții, venituri, solicitări"),
    ], [2200, 2600, 4560])
    add_callout(doc, "Problema de fond", "În mod obișnuit, magazinul online și CRM-ul sunt două programe diferite, conectate prin integrări care se pot defecta sau pot produce date neconcordante.")

    add_heading(doc, "Soluția propusă", 2)
    add_body(doc, "TrainingIT este o aplicație web care reunește cele două preocupări într-un singur produs. La suprafață funcționează ca o piață modernă de cursuri IT: utilizatorul își creează cont, consultă catalogul, cumpără, scrie recenzii și rezervă ședințe individuale cu traineri. În spate, aceleași acțiuni alimentează automat un sistem CRM folosit de echipa companiei.")
    add_bullets(doc, [
        "Înregistrarea creează un contact și un lead în CRM.",
        "Cumpărarea unui curs creează înscrierea și o afișează în istoricul administrativ.",
        "Rezervarea unei ședințe individuale poate genera automat o factură.",
        "Raportarea unei probleme creează imediat o sesizare pentru administratori.",
        "Evenimentele importante pot actualiza scoruri, statistici și jurnalul de audit.",
    ])

    add_heading(doc, "Obiectivul principal", 2)
    add_body(doc, "Obiectivul lucrării este proiectarea și implementarea unei aplicații web apropiate de forma unui produs real, capabilă să comercializeze cursuri de programare și să mențină transparent o evidență CRM completă a activității clienților.")
    add_heading(doc, "Obiective concrete", 3)
    add_bullets(doc, [
        "realizarea unui portal public pentru catalog, cumpărare, recenzii și programarea ședințelor individuale;",
        "realizarea unui portal administrativ pentru contacte, cursuri, achiziții, facturi, analize, angajați și probleme;",
        "organizarea logicii serverului în jurul șabloanelor de proiectare și al unui nucleu bazat pe evenimente;",
        "integrarea opțională a unui asistent AI pentru recomandări și traducere, fără blocarea aplicației când serviciul lipsește;",
        "stocarea persistentă a datelor operaționale și CRM într-o bază de date relațională.",
    ])
    add_callout(doc, "Cum explic oral", "Lucrarea nu prezintă doar un magazin de cursuri și nici doar un CRM. Contribuția principală este unificarea lor, astfel încât aceeași acțiune să fie văzută imediat atât de client, cât și de companie.", fill=PALE_TEAL, accent=TEAL)


def chapter_2(doc):
    add_heading(doc, "Capitolul 2. Problema abordată și soluțiile similare", 1)
    add_heading(doc, "Așteptările cursantului", 2)
    add_body(doc, "Un potențial cursant se așteaptă la experiența simplă a unui magazin online modern. El trebuie să poată înțelege rapid ce cursuri există, cât costă, cui se adresează și cum au fost evaluate de alți participanți. Procesul de cumpărare și programare trebuie să fie clar, iar raportarea unei probleme trebuie să fie ușoară.")
    add_bullets(doc, [
        "navigare clară și informații ușor de comparat;",
        "cont unic și cumpărare rapidă;",
        "recenzii credibile, permise numai după achiziție;",
        "calendar corect pentru ședințele individuale;",
        "canal simplu pentru solicitarea ajutorului.",
    ])

    add_heading(doc, "Nevoile companiei", 2)
    add_body(doc, "În spatele magazinului, compania trebuie să organizeze o activitate comercială completă. Ea urmărește leadurile și apropierea lor de cumpărare, păstrează istoricul tranzacțiilor, calculează reducerile angajaților companiilor partenere, emite facturi și exportă date pentru contabilitate.")
    add_bullets(doc, [
        "evaluarea leadurilor prin stare și scor;",
        "urmărirea achizițiilor și a veniturilor;",
        "analiza leadurilor pierdute și a înscrierilor abandonate;",
        "măsurarea interesului pentru cursuri prin accesări și rate de conversie;",
        "administrarea clienților corporate și a echipelor lor de angajați.",
    ])

    add_heading(doc, "Riscul sistemelor separate", 2)
    add_body(doc, "Atunci când magazinul și CRM-ul sunt aplicații diferite, fiecare punct de sincronizare poate produce pierderi sau diferențe de date. O achiziție poate exista în magazin, dar nu și în CRM; reducerea poate apărea pe factură, dar nu în istoricul comercial; o sesizare poate rămâne izolată de echipa care gestionează relația cu clientul.")
    add_callout(doc, "Consecință", "Reconcilierea manuală este lentă, repetitivă și predispusă la erori. TrainingIT elimină această reconciliere printr-un model comun de date.")

    add_heading(doc, "Abordări similare", 2)
    rows = [
        ("Udemy / Coursera", "Catalog, recenzii și cumpărare bine dezvoltate", "Sunt agregatoare; furnizorul individual nu deține integral relația și nu are propriul CRM integrat"),
        ("Moodle", "Livrarea și urmărirea conținutului educațional", "Comerțul electronic și CRM-ul sunt funcții secundare, adăugate prin extensii"),
        ("Salesforce / HubSpot", "Contacte, pipeline, scoruri și raportare", "Nu au implicit catalog de cursuri, calendar de traineri sau magazin public"),
        ("WooCommerce", "Poate vinde orice produs, inclusiv cursuri", "Tratează cursul ca produs obișnuit și nu înțelege leadurile, tutoratul sau instruirea corporate"),
        ("TrainingIT", "Magazin și CRM pe același backend și model de date", "Soluție specializată pentru cursuri IT, ședințe, reduceri, statistici și recomandări"),
    ]
    add_table(doc, ["Categorie", "Punct forte", "Limită / diferență"], rows, [1900, 3100, 4360])
    add_callout(doc, "Platformă agregatoare", "O platformă care reunește ofertele mai multor furnizori și intermediază relația cu utilizatorii. Furnizorul nu controlează întotdeauna integral datele, comunicarea și procesele comerciale.", fill=PALE_GRAY, accent=NAVY)

    add_heading(doc, "Delimitarea rolurilor", 2)
    add_body(doc, "TrainingIT este o aplicație unitară cu două portaluri. Rolurile USER și ADMIN sunt separate, iar accesarea unei rute rezervate celuilalt rol conduce la redirecționarea utilizatorului către propriul portal.")
    add_table(doc, ["Rol", "Cine îl folosește", "Acces principal"], [
        ("USER", "cursant individual sau angajat al unei companii", "catalog, cursuri cumpărate, recenzii, ședințe și sesizări"),
        ("ADMIN", "membru al echipei TrainingIT", "contacte, cursuri, achiziții, facturi, analize, angajați și probleme"),
    ], [1500, 3300, 4560], header_fill=PALE_TEAL)


def chapter_5(doc):
    add_heading(doc, "Capitolul 5. Arhitectura sistemului", 1)
    add_body(doc, "Aplicația este organizată în straturi care comunică prin API-uri REST și printr-un flux de evenimente în timp real. Un serviciu AI extern poate fi conectat opțional. Separarea responsabilităților face sistemul mai ușor de înțeles, testat și extins.")
    add_table(doc, ["Strat", "Responsabilitate", "Elemente principale"], [
        ("Front-end", "interfața din browser și experiența celor două portaluri", "Next.js, React, sesiune, protecția rutelor, client REST, SSE, temă și traducere"),
        ("Back-end web", "primește cereri, validează date și deleagă operațiile", "Spring Boot, controllere REST, CrmFacade, servicii pentru autentificare, AI, traineri și rapoarte"),
        ("Domeniu", "aplică regulile de funcționare ale aplicației", "comenzi, servicii, EventBus, observatori, repository-uri și DAO-uri"),
        ("Persistență", "stochează datele operaționale și CRM", "MariaDB, JDBC, HikariCP și schemă actualizată la pornire"),
        ("Serviciu extern", "oferă funcții AI opționale", "Anthropic Claude prin SDK-ul Java"),
    ], [1650, 3600, 4110])

    add_heading(doc, "Front-end", 2)
    add_body(doc, "Front-endul este o aplicație Next.js și React care afișează atât portalul clientului, cât și portalul administratorului. În browser este păstrat tipul sesiunii, USER sau ADMIN. Un mecanism de protejare a rutelor verifică accesul, iar un client REST tipizat comunică cu serverul.")
    add_bullets(doc, [
        "apelează endpointurile backendului prin cereri HTTP;",
        "se abonează la fluxul SSE pentru statisticile publice actualizate în timp real;",
        "gestionează tema luminoasă sau întunecată;",
        "oferă selectorul pentru traducerea dinamică a paginii.",
    ])

    add_heading(doc, "Back-end", 2)
    add_body(doc, "Backendul este un serviciu Spring Boot care expune controllere sub ruta /api/. Controllerele sunt intenționat subțiri: verifică datele primite și trimit execuția către nivelul de domeniu. Operațiile CRM principale sunt centralizate prin CrmFacade, în timp ce autentificarea, AI-ul, trainerii și rapoartele folosesc servicii specializate.")
    add_callout(doc, "Controller subțire", "Controllerul nu conține regulile principale ale afacerii. El transformă cererea web într-un apel către serviciul potrivit și returnează răspunsul.", fill=PALE_GRAY, accent=NAVY)

    add_heading(doc, "Nivelul de domeniu și logica de afaceri", 2)
    add_body(doc, "Nivelul de domeniu reprezintă centrul aplicației. Aici sunt implementate regulile care stabilesc cum funcționează produsul: reducerile, disponibilitatea trainerilor, crearea înscrierilor, emiterea facturilor, scorurile leadurilor și schimbările de stare.")
    add_bullets(doc, [
        "operațiile de scriere sunt încapsulate sub formă de comenzi;",
        "serviciile de domeniu coordonează regulile și entitățile;",
        "EventBus publică evenimente și notifică observatorii;",
        "repository-urile și DAO-urile asigură accesul la date.",
    ])
    add_callout(doc, "Logică de afaceri", "Totalitatea regulilor specifice domeniului care transformă o cerere într-un rezultat corect pentru companie și client. Nu înseamnă interfață și nici acces direct la baza de date.")

    add_heading(doc, "Persistența", 2)
    add_body(doc, "Toate datele sunt stocate în MariaDB. Accesul se face prin driverul JDBC și prin poolul de conexiuni HikariCP, care reutilizează eficient conexiunile. Schema cuprinde atât tabelele CRM, cât și tabelele funcționalităților web și este creată sau actualizată la pornirea aplicației.")
    add_callout(doc, "Persistență", "Datele rămân disponibile după oprirea și repornirea aplicației; ele nu sunt păstrate numai în memoria serverului.", fill=PALE_TEAL, accent=TEAL)

    add_heading(doc, "Controlul accesului", 2)
    add_steps(doc, [
        "Vizitatorul neautentificat poate deschide numai ecranele de autentificare, înregistrare și recuperare a parolei.",
        "La autentificare, adresa de e-mail este verificată în evidențele administratorilor, contactelor și angajaților.",
        "Sistemul creează o sesiune USER sau ADMIN și deschide portalul corespunzător.",
        "Pentru un angajat, sistemul identifică sau creează contactul și aplică reducerea companiei partenere.",
    ], 10)


def chapter_6(doc):
    add_heading(doc, "Capitolul 6. Portalurile și funcționalitățile aplicației", 1)
    add_heading(doc, "Autentificare și înregistrare", 2)
    add_body(doc, "Aplicația este protejată, iar un vizitator neautentificat ajunge numai la ecranele Log in, Register și Forgot password. Formularul de înregistrare este împărțit pe categorii: date personale, date de contact, adresă, profil de învățare, parolă și consimțământ GDPR obligatoriu.")
    add_body(doc, "Înregistrarea creează în CRM un contact de tip INDIVIDUAL, calculează scorul de lead și autentifică automat utilizatorul. La conectare, serverul verifică parola, stabilește rolul și transmite portalul de destinație.")

    add_heading(doc, "Portalul clientului", 2)
    add_heading(doc, "Catalogul de cursuri", 3)
    add_body(doc, "Pagina principală a portalului afișează un salut personalizat și patru valori actualizate prin SSE: numărul de cursuri, numărul de cursanți, evaluarea medie și disponibilitatea mentoratului individual. Catalogul este o grilă filtrabilă de carduri.")
    add_bullets(doc, [
        "cardul poate fi extins pentru descrierea completă;",
        "sunt afișate evaluarea medie și numărul recenziilor;",
        "utilizatorul poate deschide pagina recenziilor;",
        "butonul Enroll permite cumpărarea rapidă;",
        "prima extindere a cardului este numărată pentru analiza ratei de accesare;",
        "un chestionar AI opțional recomandă cursuri după interese, nivel și obiectiv.",
    ])

    add_heading(doc, "My Courses și recenziile", 3)
    add_body(doc, "Pagina My Courses este biblioteca personală a cursantului și afișează numai cursurile cumpărate de acesta. Recenziile pot fi scrise exclusiv aici, ceea ce leagă natural recenzia de o achiziție reală. Pentru fiecare curs cumpărat există o singură evaluare de la 1 la 5 și un comentariu; retrimiterea actualizează recenzia anterioară.")

    add_heading(doc, "Rezervarea ședințelor individuale", 3)
    add_steps(doc, [
        "Alegerea trainerului.",
        "Alegerea zilei disponibile.",
        "Alegerea orei de început și a duratei.",
        "Confirmarea rezervării și efectuarea plății.",
    ], 20)
    add_body(doc, "Trainerii lucrează de luni până sâmbătă, între 08:00 și 20:00. Zilele trecute, duminicile și zilele complet ocupate sunt dezactivate. Tariful este de 10 dolari pe oră, iar angajații partenerilor primesc automat reducerea configurată, afișată împreună cu prețul inițial.")
    add_bullets(doc, [
        "confirmarea rezervă intervalul și generează factura;",
        "intervalul devine indisponibil pentru ceilalți utilizatori;",
        "anularea din My Sessions eliberează intervalul.",
    ])

    add_heading(doc, "Asistent și raportarea problemelor", 3)
    add_body(doc, "Butonul rotund din colțul dreapta-jos deschide asistentul conversațional Claude. Butonul cu steag din stânga-jos deschide formularul Report an issue. Sesizarea ajunge imediat în secțiunea Issues, împreună cu numele și e-mailul utilizatorului.")

    add_heading(doc, "Portalul de administrare", 2)
    rows = [
        ("Contacts", "persoane și companii, stări NEW–LOST, scor automat, căutare și export CSV/Excel/PDF"),
        ("Courses", "vizualizarea internă, tabelară, a catalogului public"),
        ("Purchases", "istoricul înscrierilor, cursantul, cursul, data, evaluarea și totalul comenzilor"),
        ("Invoices", "facturile ședințelor individuale, inclusiv reducerile, cu descărcare PDF"),
        ("Analytics", "indicatori comerciali, interpretări, accesări pe curs și grafice demografice"),
        ("Employees", "angajați corporate, import Excel cu erori pe rând și recomandări AI pentru echipă"),
        ("Issues", "sesizări care pot fi rezolvate, redeschise sau șterse după confirmare explicită"),
    ]
    add_table(doc, ["Secțiune", "Rol"], rows, [1800, 7560], header_fill=PALE_TEAL)

    add_heading(doc, "Funcții disponibile pe toate paginile", 2)
    add_bullets(doc, [
        "Selectorul de traducere folosește Claude pentru conținutul static și dinamic, păstrând numele, e-mailurile și numerele.",
        "Comutatorul light/dark schimbă instantaneu tema și memorează local alegerea, aplicând-o înainte de prima afișare pentru a evita clipirea temei greșite.",
    ])
    add_callout(doc, "Interfață intuitivă", "O interfață clară și ușor de folosit de persoane fără pregătire tehnică. Termenul nu demonstrează automat conformitatea cu standardele de accesibilitate precum WCAG.", fill=PALE_GRAY, accent=NAVY)


def chapter_7(doc):
    add_heading(doc, "Capitolul 7. Tehnologiile utilizate", 1)
    add_body(doc, "TrainingIT este o aplicație full-stack deoarece include interfața din browser, serverul care implementează regulile, baza de date și mecanismele de comunicare dintre aceste componente.")
    add_table(doc, ["Zonă", "Tehnologie", "Motivul utilizării"], [
        ("Front-end", "Next.js 16 și React 19", "componente reutilizabile și experiență unitară pentru ambele portaluri"),
        ("Front-end", "TypeScript 5", "tipizare statică și identificarea timpurie a erorilor"),
        ("Stilizare", "Tailwind CSS 4", "temă vizuală personalizată și mod întunecat"),
        ("Back-end", "Java 17", "implementarea nivelului web și a domeniului orientat pe obiecte"),
        ("Back-end", "Spring Boot 3.5", "configurare, REST, SSE și conectarea componentelor"),
        ("Validare", "Spring Validation", "reguli declarative pentru payloadurile primite"),
        ("Cod Java", "Lombok", "reducerea codului repetitiv în entități și DTO-uri"),
        ("Logging", "SLF4J și Logback", "mesaje, avertismente și erori"),
        ("Bază de date", "MariaDB", "stocarea relațională a datelor operaționale și CRM"),
        ("Acces la date", "JDBC și HikariCP", "SQL explicit și conexiuni reutilizate"),
        ("Documente", "Apache POI și OpenPDF", "Excel, importuri, exporturi și facturi PDF"),
        ("AI", "Anthropic Java SDK", "asistent, recomandări și traducere, activate numai cu o cheie API"),
    ], [1500, 2450, 5410])

    add_heading(doc, "De ce nu este folosit un ORM", 2)
    add_body(doc, "Nivelul de domeniu folosește repository-uri și DAO-uri scrise manual. Această alegere păstrează SQL-ul explicit și face vizibile șabloanele de proiectare folosite pentru accesul la date. Dezavantajul este că dezvoltatorul trebuie să scrie și să întrețină mai mult cod de infrastructură.")
    add_callout(doc, "Full-stack", "Un sistem care acoperă întregul traseu al datelor: interfață → API → logică de domeniu → bază de date → răspuns către utilizator.", fill=PALE_TEAL, accent=TEAL)


def chapter_8(doc):
    add_heading(doc, "Capitolul 8. Implementarea și modelul de date", 1)
    add_heading(doc, "Organizarea codului", 2)
    add_body(doc, "Codul serverului conține aproximativ 11.500 de linii Java, organizate pe responsabilități: builders, commands, DAOs, factories, models, observers, patterns, repositories, services, strategies, validation și web. Clientul cuprinde aproximativ cincizeci de module TypeScript și React.")

    add_heading(doc, "Nucleul bazat pe evenimente", 2)
    add_body(doc, "Comportamentul „o acțiune, mai multe reacții” este realizat prin șablonul Observer și EventBus. După finalizarea unei operații importante, serviciul publică un eveniment. EventBus notifică fiecare observator înregistrat, iar observatorii reacționează independent.")
    add_table(doc, ["Eveniment", "Reacții posibile"], [
        ("Contact creat", "audit, mesaj de bun venit și calcularea scorului inițial"),
        ("Înscriere creată", "audit, confirmare, actualizarea statisticilor și a stării leadului"),
        ("Ședință rezervată", "generarea facturii și confirmarea rezervării"),
        ("Activitate finalizată", "actualizarea procesului comercial"),
        ("Stare lead / oportunitate schimbată", "audit și actualizarea evidențelor asociate"),
    ], [2900, 6460], header_fill=PALE_TEAL)
    add_callout(doc, "Avantaj", "Serviciul care face rezervarea nu trebuie să cunoască direct toate reacțiile ulterioare. Observatorii pot fi adăugați sau modificați fără a încărca operația principală.")

    add_heading(doc, "Idempotență și izolarea erorilor", 2)
    add_body(doc, "Crearea facturii este idempotentă: dacă există deja o factură pentru ședință, sistemul nu generează una nouă. În plus, eroarea observatorului este izolată. O problemă la facturare este înregistrată în log, dar nu anulează rezervarea care a produs evenimentul.")
    add_table(doc, ["Proprietate", "Ce previne", "Exemplu"], [
        ("Idempotență", "efecte duplicate", "același eveniment nu produce două facturi"),
        ("Izolarea erorii", "propagarea eșecului secundar", "rezervarea rămâne salvată chiar dacă factura eșuează"),
        ("Degradare controlată", "blocarea aplicației de o funcție opțională", "lipsa cheii AI ascunde butoanele, restul continuă"),
    ], [2200, 3300, 3860])

    add_heading(doc, "Modelul de date", 2)
    add_body(doc, "Modelul comun include atât entitățile CRM, cât și entitățile funcționalităților web. Contact poate fi o persoană INDIVIDUAL sau o companie CORPORATE. Companiile au Employee, iar cursurile oferă CourseSession. Enrollment asociază un contact cu o sesiune de curs.")
    add_bullets(doc, [
        "Opportunity și Activity modelează procesul comercial;",
        "Trainer și MediationSession modelează ședințele individuale;",
        "ședințele individuale pot genera SessionInvoice;",
        "recenziile sunt stocate direct în Enrollment prin rating și feedback.",
    ])

    add_heading(doc, "De ce recenzia este în Enrollment", 2)
    add_body(doc, "Structura impune natural o singură recenzie pentru perechea contact–curs. O nouă trimitere înlocuiește evaluarea anterioară. Valorile 1–5 reprezintă evaluări valide, iar 0 înseamnă că utilizatorul nu a evaluat încă acel curs.")
    add_callout(doc, "Legătură importantă", "Recenzia nu este independentă de achiziție. Sistemul poate verifica faptul că utilizatorul a cumpărat cursul înainte de a-i permite evaluarea.", fill=PALE_TEAL, accent=TEAL)

    add_heading(doc, "Ordinea corectă contact–înscriere", 2)
    add_steps(doc, [
        "La înregistrare, utilizatorul este creat ca contact în CRM.",
        "La cumpărare, serverul caută contactul după e-mail.",
        "Dacă nu îl găsește, îl creează înainte de continuarea fluxului.",
        "Serverul identifică sau creează sesiunea cursului.",
        "Abia apoi creează Enrollment și publică evenimentul de înscriere.",
    ], 30)
    add_callout(doc, "Clarificare", "Contactul nu este creat după înscriere. Enrollment are nevoie de contact pentru a putea fi asociat corect.", fill=PALE_GRAY, accent=NAVY)

    add_heading(doc, "Integrarea AI", 2)
    add_body(doc, "ClaudeClient este un wrapper subțire peste SDK-ul oficial Anthropic. Cheia este citită din configurare. Dacă lipsește, clientul rămâne neinițializat, isEnabled() întoarce false, iar endpointurile raportează că funcția este oprită în loc să producă o eroare fatală.")
    add_bullets(doc, [
        "asistent conversațional pentru întrebări despre cursuri și înscriere;",
        "chestionar de recomandare pentru vizitatori;",
        "recomandări de formare pentru angajații unei companii;",
        "traducerea dinamică a paginii.",
    ])

    add_heading(doc, "SSE, statistici și documente", 2)
    add_body(doc, "Statisticile publice sunt transmise browserului prin Server-Sent Events. Broadcasterul recalculează valorile cât timp există cel puțin un client conectat și transmite numai modificările. Astfel, pagina se actualizează fără interogări periodice.")
    add_body(doc, "Apache POI este utilizat pentru fișiere Excel, iar OpenPDF pentru facturi și exporturi PDF. Aceleași biblioteci susțin importul angajaților și exporturile administrative.")


def chapter_9(doc):
    add_heading(doc, "Capitolul 9. Testarea aplicației", 1)
    add_heading(doc, "Strategia de testare", 2)
    add_body(doc, "Corectitudinea TrainingIT depinde mai ales de propagarea completă a unei acțiuni prin sistem. De aceea, testarea s-a concentrat pe validarea datelor, robustețea nucleului bazat pe evenimente și verificarea funcțională a fluxurilor end-to-end.")

    add_heading(doc, "Validarea pe două niveluri", 2)
    add_table(doc, ["Nivel", "Ce verifică", "Scop"], [
        ("Client", "câmpuri lipsă, confirmarea parolei și consimțământul GDPR", "feedback rapid înainte de trimiterea formularului"),
        ("Server", "regulile Spring Validation pentru payload", "protecția API-ului indiferent de client"),
        ("Domeniu", "lanțul Chain of Responsibility pentru e-mail, telefon, GDPR și tipul contactului", "oprirea datelor invalide înainte de baza de date"),
    ], [1500, 4600, 3260])
    add_callout(doc, "Spring Validation", "Mecanism declarativ prin care adnotările descriu cerințele datelor primite, iar Spring respinge automat cererile invalide.", fill=PALE_TEAL, accent=TEAL)
    add_callout(doc, "Chain of Responsibility", "Șablon în care datele trec succesiv prin mai mulți verificatori. Fiecare componentă răspunde de o singură regulă și poate opri lanțul când găsește o eroare.", fill=PALE_GRAY, accent=NAVY)

    add_heading(doc, "Robustețea nucleului bazat pe evenimente", 2)
    add_body(doc, "Robustețea înseamnă capacitatea sistemului de a rămâne stabil și previzibil atunci când apar erori, evenimente repetate sau servicii indisponibile. Au fost verificate explicit următoarele comportamente:")
    add_bullets(doc, [
        "evenimentul repetat nu generează facturi duplicate;",
        "eroarea observatorului de facturare nu anulează rezervarea;",
        "lipsa cheii API dezactivează funcțiile AI fără a afecta restul aplicației;",
        "ștergerea unei sesizări necesită confirmarea explicită a acțiunii ireversibile.",
    ])

    add_heading(doc, "Fluxurile testate manual", 2)
    tests = [
        ("Înregistrare", "crearea contactului și a leadului cu scor"),
        ("Autentificare", "direcționarea USER și ADMIN către portalul corect"),
        ("Înscriere", "apariția imediată în Purchases"),
        ("Recenzie", "actualizarea evaluării medii a cursului"),
        ("Ședință individuală", "factură și reducere corectă pentru angajat"),
        ("Sesizare", "apariția în Issues"),
        ("Export", "generarea fișierelor CSV, Excel și PDF"),
    ]
    add_table(doc, ["Flux", "Rezultat verificat"], tests, [2800, 6560])
    add_body(doc, "Testele au fost executate pe o instanță funcțională conectată la o bază de date MariaDB reală. Proiectul include spring-boot-starter-test ca bază pentru o suită automată viitoare, însă acoperirea completă cu teste unitare și de integrare rămâne o direcție de dezvoltare.")
    add_callout(doc, "Limitare formulată corect", "Aplicația a fost verificată funcțional, dar nu trebuie afirmat că există deja o acoperire automată completă. Aceasta este planificată pentru viitor.", fill=PALE_GRAY, accent=GOLD)


def chapter_10(doc):
    add_heading(doc, "Capitolul 10. Scenarii de utilizare", 1)
    add_heading(doc, "Alexandru – cursant individual", 2)
    add_body(doc, "Alexandru lucrează în vânzări și dorește să treacă în IT, dar nu știe de unde să înceapă. Asistentul îi recomandă două cursuri pentru începători. El se înregistrează, apare în CRM ca lead cu scor calculat, cumpără cursul mai accesibil și rezervă o ședință individuală. Factura este produsă automat.")
    add_callout(doc, "Rezultat", "Un vizitator nehotărât devine client plătitor în câteva minute, iar compania obține automat contactul, leadul, achiziția și factura.", fill=PALE_TEAL, accent=TEAL)

    add_heading(doc, "Cristina – client venit prin recomandare", 2)
    add_body(doc, "Cristina este dezvoltator junior și ajunge prin recomandarea unui prieten, ceea ce contribuie la un scor ridicat. Ea solicită direct cursuri avansate, află despre ședințele individuale și cumpără rapid. Scenariul ilustrează valoarea comercială a recomandărilor și costul redus de achiziție al acestor clienți.")
    add_callout(doc, "Rezultat", "Sursa leadului poate fi analizată în CRM, iar compania poate observa că recomandările clienților mulțumiți sunt eficiente și ieftine.")

    add_heading(doc, "BankTech Solutions – client corporate", 2)
    add_body(doc, "Administratorul importă zeci de angajați dintr-un fișier Excel și solicită recomandări AI pentru întreaga echipă. În CRM este creată o oportunitate B2B. Angajații cumpără cu reducerea automată de 60%, iar fiecare ședință individuală generează factura corespunzătoare.")
    add_callout(doc, "Rezultat", "O singură companie aduce simultan numeroși cursanți, iar sistemul urmărește echipa, reducerile și facturile fără introducere manuală repetată.", fill=PALE_TEAL, accent=TEAL)

    add_heading(doc, "Fluxul comun de cumpărare și recenzie", 2)
    add_steps(doc, [
        "Identificarea sau crearea contactului.",
        "Identificarea sau crearea sesiunii cursului.",
        "Salvarea Enrollment.",
        "Publicarea evenimentului de înscriere.",
        "Recalcularea statisticilor și trimiterea valorilor modificate prin SSE.",
        "La recenzie, verificarea achiziției anterioare.",
        "Salvarea evaluării și comentariului în Enrollment.",
    ], 40)


def chapter_11(doc):
    add_heading(doc, "Capitolul 11. Concluzii și dezvoltări viitoare", 1)
    add_heading(doc, "Rezultatul lucrării", 2)
    add_body(doc, "TrainingIT demonstrează că un magazin de cursuri și un CRM pot fi realizate ca un singur produs coerent. Cele două portaluri folosesc același backend și același model de date, astfel încât fiecare acțiune este înregistrată o singură dată și devine imediat disponibilă proceselor administrative relevante.")
    add_body(doc, "Vizitatorul devine lead la înregistrare și client la prima achiziție. Cumpărările, recenziile, ședințele și sesizările sunt reflectate automat în portalul administratorului. Nucleul bazat pe evenimente permite reacții multiple fără cuplarea directă a fiecărei operații de toate efectele sale secundare.")

    add_heading(doc, "Îndeplinirea obiectivelor", 2)
    add_bullets(doc, [
        "portal public complet pentru descoperirea și cumpărarea cursurilor;",
        "portal administrativ organizat în șapte secțiuni;",
        "nivel de domeniu bazat pe șabloane clasice și evenimente;",
        "persistența datelor în MariaDB;",
        "funcții AI și traducere disponibile opțional;",
        "comportament controlat când serviciile opționale lipsesc.",
    ])
    add_callout(doc, "Caracter distinctiv", "TrainingIT combină o interfață ușor de folosit, un nucleu CRM bine structurat și un nivel opțional de inteligență artificială într-un singur produs.", fill=PALE_TEAL, accent=TEAL)

    add_heading(doc, "Direcții viitoare", 2)
    add_table(doc, ["Direcție", "Extindere propusă"], [
        ("Testare automată", "teste unitare și de integrare pentru comenzi, observatori, strategii, endpointuri și integrare continuă"),
        ("Autentificare", "tokenuri, parole stocate prin hashing, verificarea e-mailului și resetare prin linkuri care expiră"),
        ("Analize", "tendințe istorice, cohorte și dashboarduri exportabile"),
        ("AI", "context bazat pe catalogul live, memorie conversațională și rezumarea sesizărilor"),
        ("Implementare în cloud", "containerizarea serviciilor, configurare externă, scalare și pregătire pentru producție"),
    ], [2200, 7160])


def appendix_glossary(doc):
    add_heading(doc, "Anexa A. Glosar explicat", 1)
    glossary = [
        ("CRM", "Sistem pentru administrarea contactelor, leadurilor, vânzărilor, facturilor și interacțiunilor cu clienții."),
        ("Lead", "Persoană sau companie care poate deveni client și este urmărită prin stare și scor."),
        ("Platformă agregatoare", "Platformă care reunește ofertele mai multor furnizori și intermediază relația cu utilizatorii."),
        ("Jurnal de audit", "Evidență cronologică a acțiunilor importante: cine a făcut acțiunea, când și asupra cărei înregistrări."),
        ("Logică de afaceri", "Regulile specifice aplicației: reduceri, disponibilitate, facturare, scoruri și schimbări de stare."),
        ("Full-stack", "Sistem care include interfața, backendul, baza de date și comunicarea dintre ele."),
        ("REST API", "Interfață HTTP prin care front-endul trimite cereri backendului și primește rezultate."),
        ("SSE", "Server-Sent Events; mecanism prin care serverul transmite actualizări în timp real către browser."),
        ("DTO", "Obiect utilizat pentru transportul datelor între client, controllere și servicii."),
        ("DAO", "Componentă care execută operații directe de acces la baza de date."),
        ("Repository", "Abstracție care oferă operații de persistență potrivite domeniului aplicației."),
        ("Observer", "Șablon prin care o componentă reacționează la un eveniment fără a fi legată direct de operația care l-a produs."),
        ("EventBus", "Mecanism central care publică evenimente și notifică observatorii înregistrați."),
        ("Idempotent", "O operație repetată produce același rezultat și nu creează efecte duplicate."),
        ("Spring Validation", "Mecanism Spring pentru verificarea declarativă a datelor primite de backend."),
        ("Chain of Responsibility", "Șablon care transmite datele printr-un lanț de verificatori, fiecare cu o responsabilitate distinctă."),
        ("Robustețe", "Capacitatea sistemului de a rămâne stabil și previzibil în prezența erorilor sau situațiilor neașteptate."),
        ("Degradare controlată", "Dezactivarea unei funcții opționale fără blocarea sau afectarea celorlalte funcționalități."),
        ("Interfață intuitivă", "Interfață clară și ușor de folosit; afirmația nu echivalează automat cu certificarea accesibilității."),
        ("HikariCP", "Pool de conexiuni care reutilizează conexiunile la baza de date pentru eficiență."),
    ]
    add_table(doc, ["Termen", "Explicație în contextul TrainingIT"], glossary, [2600, 6760], header_fill=PALE_TEAL)


def appendix_oral(doc):
    add_heading(doc, "Anexa B. Răspunsuri scurte pentru susținere", 1)
    add_callout(doc, "Formulare utilă", "TrainingIT nu conectează două sisteme independente; implementează două experiențe peste același nucleu și aceeași bază de date.", fill=PALE_TEAL, accent=TEAL)
    qa = [
        ("Care este problema principală?", "Magazinul online și CRM-ul sunt de obicei separate, iar sincronizarea poate produce date lipsă sau contradictorii."),
        ("Care este contribuția aplicației?", "Unifică portalul clientului și CRM-ul pe același backend și model de date, astfel încât acțiunile nu trebuie reconciliate manual."),
        ("Ce înseamnă «o acțiune, mai multe reacții»?", "O rezervare sau o achiziție publică un eveniment, iar observatorii pot genera factura, actualiza scorul, scrie auditul și trimite confirmări."),
        ("De ce este util EventBus?", "Reduce legăturile directe dintre operația principală și reacțiile secundare, ceea ce face codul mai ușor de extins."),
        ("Ce este jurnalul de audit?", "Istoricul acțiunilor importante, util pentru trasabilitate și verificare."),
        ("De ce este idempotentă facturarea?", "Pentru ca repetarea accidentală a unui eveniment să nu producă facturi duplicate."),
        ("Ce se întâmplă dacă AI-ul nu este configurat?", "Butoanele AI dispar, endpointurile raportează funcția oprită, iar restul aplicației continuă să funcționeze."),
        ("Când este creat contactul?", "La înregistrare; iar la cumpărare este căutat sau creat înainte de Enrollment, niciodată după."),
        ("Cum sunt validate datele?", "În browser, prin Spring Validation pe server și printr-un lanț Chain of Responsibility pentru datele contactului."),
        ("Care este limita principală a testării?", "Fluxurile au fost verificate manual end-to-end, dar suita automată completă rămâne o dezvoltare viitoare."),
        ("De ce aplicația este full-stack?", "Include clientul, API-ul, logica de domeniu, persistența și comunicarea în timp real."),
        ("Ce diferențiază TrainingIT de Udemy?", "TrainingIT este produsul unui furnizor care își deține relația cu clientul și are propriul CRM integrat; Udemy este o platformă agregatoare."),
    ]
    add_table(doc, ["Întrebare", "Răspuns recomandat"], qa, [3300, 6060], header_fill=PALE_BLUE)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_running_furniture(doc.sections[0])
    doc.core_properties.title = "Ghid de recapitulare TrainingIT"
    doc.core_properties.subject = "Conspect explicat al lucrării de licență TrainingIT"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "TrainingIT, CRM, cursuri IT, lucrare de licență"

    add_numbering_definition(doc, 1, 1, fmt="bullet")

    add_cover(doc)
    add_reading_map(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    chapter_8(doc)
    chapter_9(doc)
    chapter_10(doc)
    chapter_11(doc)
    appendix_glossary(doc)
    appendix_oral(doc)

    # Keep table rows intact and ensure all table text remains readable.
    for table in doc.tables:
        for row in table.rows:
            prevent_row_split(row)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
