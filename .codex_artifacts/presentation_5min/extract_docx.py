from pathlib import Path
from docx import Document


source = Path(r"D:\TrainingIT_site\.codex_artifacts\presentation_5min\TrainingIT_conspectare.docx")
output = source.with_name("conspect_text.txt")
doc = Document(source)

lines = []
for index, paragraph in enumerate(doc.paragraphs, start=1):
    text = paragraph.text.strip()
    if text:
        style = paragraph.style.name if paragraph.style else ""
        lines.append(f"P{index:03d} [{style}] {text}")

for table_index, table in enumerate(doc.tables, start=1):
    lines.append(f"TABLE {table_index}")
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
        lines.append(" | ".join(cells))

output.write_text("\n".join(lines), encoding="utf-8")
print(output)
