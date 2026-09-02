from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader


sys.stdout.reconfigure(encoding="utf-8", errors="replace")


def clean(text: str) -> str:
    return re.sub(r"[ \t]+", " ", text.replace("\x00", "")).strip()


def inspect_pdf(path: Path, query: str | None, context: int) -> None:
    reader = PdfReader(str(path))
    print(f"PDF pages: {len(reader.pages)}")
    page_texts = [page.extract_text() or "" for page in reader.pages]
    if query:
        pattern = re.compile(query, re.IGNORECASE)
        for page_no, text in enumerate(page_texts, start=1):
            lines = [clean(line) for line in text.splitlines() if clean(line)]
            for index, line in enumerate(lines):
                if pattern.search(line):
                    start = max(0, index - context)
                    end = min(len(lines), index + context + 1)
                    print(f"\n=== PAGE {page_no}, LINE {index + 1} ===")
                    print("\n".join(lines[start:end]))
    else:
        for page_no, text in enumerate(page_texts, start=1):
            print(f"\n=== PAGE {page_no} ===")
            print(clean(text))


def inspect_docx(path: Path) -> None:
    doc = Document(str(path))
    print(f"DOCX paragraphs: {len(doc.paragraphs)}; tables: {len(doc.tables)}")
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = clean(paragraph.text)
        if text:
            print(f"P{index:03d} [{paragraph.style.name}]: {text}")
    for table_index, table in enumerate(doc.tables, start=1):
        print(f"\n=== TABLE {table_index} ===")
        for row in table.rows:
            print(" | ".join(clean(cell.text) for cell in row.cells))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("path", type=Path)
    parser.add_argument("--query")
    parser.add_argument("--context", type=int, default=3)
    args = parser.parse_args()
    if args.path.suffix.lower() == ".pdf":
        inspect_pdf(args.path, args.query, args.context)
    elif args.path.suffix.lower() == ".docx":
        inspect_docx(args.path)
    else:
        raise SystemExit("Expected PDF or DOCX")


if __name__ == "__main__":
    main()
