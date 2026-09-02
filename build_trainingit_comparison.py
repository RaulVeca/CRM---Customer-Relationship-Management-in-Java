from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\TrainingIT_site\output\documents\Sinteza_comparativa_parte_tehnica_TrainingIT.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MID_BLUE = "DCEAF6"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E7F4F3"
TEAL = "176B73"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE6"
TEXT = "202124"
MUTED = "5F6B7A"
WHITE = "FFFFFF"
GREEN = "2E6B4F"
GREEN_FILL = "E7F2EC"
GOLD = "8A6500"
GOLD_FILL = "FFF4D6"
RED = "8F2D2D"
RED_FILL = "FBEAEA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    assert sum(widths_dxa) == 9360, widths_dxa
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    for tag_name, value in (("tblW", 9360), ("tblInd", indent_dxa)):
        tag = tbl_pr.find(qn(f"w:{tag_name}"))
        if tag is None:
            tag = OxmlElement(f"w:{tag_name}")
            tbl_pr.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color=BLUE, size=12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def setup_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def add_abstract(abs_id: int, num_fmt: str, lvl_text: str, left: int, hanging: int, font=None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    add_abstract(next_abs, "bullet", "•", 540, 270, "Calibri")
    add_abstract(next_abs + 1, "decimal", "%1.", 540, 270)

    def add_num(num_id: int, abs_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_id = OxmlElement("w:abstractNumId")
        abstract_id.set(qn("w:val"), str(abs_id))
        num.append(abstract_id)
        numbering.append(num)

    add_num(next_num, next_abs)
    add_num(next_num + 1, next_abs + 1)
    return next_num, next_num + 1


def set_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_bullet(doc, text, bullet_id, bold_prefix=None):
    p = doc.add_paragraph()
    set_num(p, bullet_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=TEXT)
    else:
        r = p.add_run(text)
        set_run_font(r, color=TEXT)
    return p


def add_number(doc, text, number_id):
    p = doc.add_paragraph()
    set_num(p, number_id)
    r = p.add_run(text)
    set_run_font(r, color=TEXT)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=BLUE, after=10):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color=fill, size=2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.20
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=label_color)
    r = p.add_run(text)
    set_run_font(r, color=TEXT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(after)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_source_line(doc, first_pages, second_pages):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    keep_with_next(p)
    r = p.add_run("Repere: ")
    set_run_font(r, size=9.5, bold=True, color=MUTED)
    r = p.add_run(f"documentul 1 - {first_pages}; documentul 2 - {second_pages}.")
    set_run_font(r, size=9.5, italic=True, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=TEXT)
    else:
        r = p.add_run(text)
        set_run_font(r, color=TEXT)
    return p


def add_flow_step(doc, label, detail, fill):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shade_paragraph(p, fill)
    r = p.add_run(label)
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(f"  {detail}")
    set_run_font(r, color=TEXT)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, MUTED, 0, 18),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        p = header.paragraphs[0]
        p.text = "TRAININGIT  |  SINTEZĂ TEHNICĂ COMPARATIVĂ"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_run_font(run, size=8.5, bold=True, color=MUTED)
        add_bottom_border(p, color=MID_GRAY, size=4)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(p)


def style_table_text(table, header=True, body_size=9.2) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.12
                for run in p.runs:
                    set_run_font(
                        run,
                        size=9.5 if header and row_idx == 0 else body_size,
                        bold=True if header and row_idx == 0 else None,
                        color=NAVY if header and row_idx == 0 else TEXT,
                    )


def add_summary_matrix(doc):
    rows = [
        ("Arhitectură full-stack și REST/JSON", "DIRECTĂ", "D1: p. 1, 5-7, 15 | D2: p. 1-2, 7-9, 26-28, 33, 46-47"),
        ("Stack tehnologic", "DIRECTĂ", "D1: p. 1, 9-10 | D2: p. 1-2, 8, 26, 33, 46"),
        ("Straturi backend și CrmFacade", "DIRECTĂ", "D1: p. 5-7 | D2: p. 26-28, 41-44"),
        ("MariaDB, JDBC, HikariCP, fără ORM", "DIRECTĂ", "D1: p. 7, 9-10 | D2: p. 1-2, 8, 26-32, 42-43, 46-48"),
        ("Șabloane de proiectare", "DIRECTĂ", "D1: p. 6-7, 10-13 | D2: p. 7, 27-28, 38, 41-45, 47"),
        ("EventBus / Observer și reacții multiple", "DIRECTĂ", "D1: p. 7, 10-11, 15 | D2: p. 27-28, 38, 43-45, 47"),
        ("Model de date comun", "DIRECTĂ", "D1: p. 10-12 | D2: p. 29-31, 38-44"),
        ("Flux cumpărare / înscriere", "DIRECTĂ", "D1: p. 3, 10-12, 14-15 | D2: p. 18, 20, 38-45"),
        ("Frontend, SSE, temă și traducere", "DIRECTĂ", "D1: p. 6, 8-9, 12 | D2: p. 18, 33-37, 43, 46-47"),
        ("AI opțional și degradare controlată", "DIRECTĂ", "D1: p. 1, 9, 12, 15 | D2: p. 8, 18-22, 32, 36, 47"),
        ("Portaluri și funcționalități", "CONTEXTUALĂ", "D1: p. 7-9 | D2: p. 18-25, 34-35"),
        ("Autentificare și control acces", "PARȚIALĂ", "D1: p. 7, 9 | D2: p. 18-20, 32, 35, 46-48"),
        ("Plăți și notificări externe", "PARȚIALĂ", "D1: p. 3, 8-11 | D2: p. 18, 21-22, 38-45, 47-48"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3150, 1650, 4560])
    set_table_borders(table)
    headers = ["Temă tehnică", "Nivel", "Repere verificate"]
    for i, value in enumerate(headers):
        table.cell(0, i).text = value
        set_cell_shading(table.cell(0, i), PALE_BLUE)
    repeat_header(table.rows[0])
    for topic, status, refs in rows:
        cells = table.add_row().cells
        cells[0].text = topic
        cells[1].text = status
        cells[2].text = refs
        fill = GREEN_FILL if status == "DIRECTĂ" else (PALE_TEAL if status == "CONTEXTUALĂ" else GOLD_FILL)
        set_cell_shading(cells[1], fill)
        for p in cells[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, [3150, 1650, 4560])
    style_table_text(table, body_size=8.7)
    for row in table.rows[1:]:
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=8.4, bold=True, color=GREEN if row.cells[1].text == "DIRECTĂ" else GOLD)
    return table


def add_technology_table(doc):
    rows = [
        ("Frontend", "Next.js 16, React 19, TypeScript 5, Tailwind CSS v4", "Confirmare explicită în ambele documente."),
        ("Backend", "Java + Spring Boot 3.5", "Java este confirmat; versiunea Java 17 apare în conspect, dar nu este precizată explicit în PDF."),
        ("Comunicare", "REST/JSON și Server-Sent Events (SSE)", "REST pentru cereri/răspunsuri; SSE pentru actualizarea statisticilor publice."),
        ("Persistență", "MariaDB, JDBC nativ, HikariCP", "SQL explicit, pool de conexiuni și schemă inițializată de aplicație."),
        ("AI", "SDK Anthropic / Claude", "Asistent, recomandări și traducere; modul opțional."),
        ("Documente", "Apache POI și OpenPDF", "Import/export Excel și generare/export PDF."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1650, 3250, 4460])
    set_table_borders(table)
    for i, text in enumerate(("Zonă", "Tehnologie comună", "Clarificare")):
        table.cell(0, i).text = text
        set_cell_shading(table.cell(0, i), PALE_BLUE)
    repeat_header(table.rows[0])
    for a, b, c in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = a, b, c
    set_table_geometry(table, [1650, 3250, 4460])
    style_table_text(table, body_size=9.0)


def add_patterns_table(doc):
    rows = [
        ("Facade", "CrmFacade este punctul unic de intrare în domeniu și ascunde complexitatea serviciilor."),
        ("Singleton", "Servicii și componente centrale sunt accesate ca instanțe unice."),
        ("Command", "Operațiile de modificare sunt încapsulate; apare mai ales în procesele administrative."),
        ("Repository + DAO", "Repository oferă interfața de domeniu; DAO execută SQL-ul JDBC către MariaDB."),
        ("Builder", "Construiește obiecte complexe, inclusiv Enrollment."),
        ("Observer / EventBus", "Evenimentele sunt publicate central și consumate de observatori independenți."),
        ("Factory", "Construiește notificări pentru canale diferite, precum e-mail și SMS."),
        ("Strategy", "Alege algoritmul de scorare a leadului în funcție de tipul contactului."),
        ("Chain of Responsibility", "Validează succesiv e-mailul, telefonul, tipul contactului și consimțământul GDPR."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2400, 6960])
    set_table_borders(table)
    for i, text in enumerate(("Șablon", "Rol confirmat în lucrarea a doua")):
        table.cell(0, i).text = text
        set_cell_shading(table.cell(0, i), PALE_BLUE)
    repeat_header(table.rows[0])
    for name, role in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = name, role
    set_table_geometry(table, [2400, 6960])
    style_table_text(table, body_size=9.2)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_id, number_id = setup_numbering(doc)

    # Cover - editorial_cover pattern, adapted to a technical comparison report.
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("RAPORT TEHNIC COMPARATIV")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Partea tehnică comună")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("dintre conspectul TrainingIT și lucrarea de licență")
    set_run_font(r, size=16, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("Identificare • sistematizare • clarificare • trasabilitate pe pagini")
    set_run_font(r, size=10.5, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documentul 1: TrainingIT_conspectare.docx\nDocumentul 2: newfile4.pdf")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("31 august 2026")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "1. Concluzia comparării", 1)
    add_callout(
        doc,
        "Verdict",
        "Cea mai mare parte a nucleului tehnic din conspect se regăsește explicit în lucrarea de licență. "
        "Potrivirea este foarte puternică pentru arhitectura full-stack, stackul tehnologic, separarea pe straturi, "
        "șabloanele de proiectare, persistență, modelul de date și fluxul de cumpărare/înscriere.",
        fill=GREEN_FILL,
        label_color=GREEN,
    )
    add_body(
        doc,
        "Documentul 1 este o reformulare didactică și mai accesibilă a unor idei tehnice din documentul 2. "
        "El comprimă explicațiile și, în anumite locuri, generalizează implementarea. De aceea, potrivirea trebuie "
        "citită pe trei niveluri: directă, contextuală și parțială cu precizări.",
    )

    add_heading(doc, "Cum se interpretează nivelurile", 2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 2520, 5040])
    set_table_borders(table)
    for i, text in enumerate(("Nivel", "Semnificație", "Cum trebuie folosit")):
        table.cell(0, i).text = text
        set_cell_shading(table.cell(0, i), PALE_BLUE)
    repeat_header(table.rows[0])
    legend = [
        ("DIRECTĂ", "Aceeași tehnologie, structură sau funcție este menționată explicit.", "Poate fi prezentată ca parte tehnică confirmată."),
        ("CONTEXTUALĂ", "Ideea este aceeași, dar lucrarea a doua o descrie mai larg sau cu altă terminologie.", "Se păstrează ideea și se adaugă explicația din lucrarea a doua."),
        ("PARȚIALĂ", "Există un nucleu comun, dar și o limită ori o diferență importantă.", "Se folosește numai împreună cu precizarea indicată."),
    ]
    for level, meaning, use in legend:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = level, meaning, use
        set_cell_shading(cells[0], GREEN_FILL if level == "DIRECTĂ" else (PALE_TEAL if level == "CONTEXTUALĂ" else GOLD_FILL))
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, [1800, 2520, 5040])
    style_table_text(table, body_size=9.1)

    add_heading(doc, "2. Matrice rapidă de corespondență", 1)
    add_body(doc, "Tabelul următor arată, dintr-o privire, ce elemente tehnice din primul document sunt susținute de al doilea document.")
    add_summary_matrix(doc)

    add_heading(doc, "3. Arhitectura generală și traseul datelor", 1)
    add_source_line(doc, "p. 1, 5-7, 9, 15", "p. 1-2, 7-9, 26-28, 33, 41-47")
    add_body(
        doc,
        "Ambele documente descriu aceeași soluție full-stack: o interfață Next.js separată de un backend Java/Spring Boot, "
        "cu schimb de date printr-un contract REST/JSON. Backendul păstrează logica de afaceri, schema relațională și accesul la date; "
        "frontendul gestionează experiența utilizatorului, rutarea și starea locală.",
    )
    add_heading(doc, "Flux tehnic, în ordine", 2)
    add_flow_step(doc, "1. Frontend", "Next.js 16 + React 19 + TypeScript 5 + Tailwind CSS v4", PALE_BLUE)
    add_flow_step(doc, "2. Contract", "cereri HTTP prin REST/JSON; actualizări live prin SSE", LIGHT_GRAY)
    add_flow_step(doc, "3. Strat web", "controller Spring Boot subțire, fără logică principală de afaceri", PALE_BLUE)
    add_flow_step(doc, "4. Domeniu", "CrmFacade → servicii → comenzi / repository-uri", PALE_TEAL)
    add_flow_step(doc, "5. Persistență", "DAO → JDBC → HikariCP → MariaDB", LIGHT_GRAY)
    add_flow_step(doc, "6. Reacții laterale", "EventBus / Observer pentru audit, notificări și actualizări de stare", PALE_BLUE)
    add_flow_step(doc, "7. Serviciu extern", "Claude API, activat numai când există configurarea necesară", PALE_TEAL)
    add_callout(
        doc,
        "Clarificare importantă",
        "Conspectul vorbește despre «o singură aplicație». Lucrarea de licență precizează mai exact că produsul este format din "
        "două subsisteme autonome - frontend și backend - care aparțin aceleiași soluții și comunică strict prin REST/JSON. "
        "Așadar, este un singur produs funcțional, nu un monolit tehnic.",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    add_heading(doc, "4. Tehnologiile confirmate în ambele documente", 1)
    add_source_line(doc, "p. 1, 7, 9-12", "p. 1-2, 8-9, 26, 32-37, 46")
    add_technology_table(doc)
    add_callout(
        doc,
        "Precizie",
        "Versiunea Java 17, Lombok, SLF4J/Logback și Spring Validation apar în conspect, dar nu sunt confirmate explicit în PDF. "
        "Din comparație pot fi susținute sigur Java, Spring Boot 3.5, MariaDB, JDBC și HikariCP.",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    add_heading(doc, "5. Arhitectura backend pe straturi", 1)
    add_source_line(doc, "p. 5-7, 10", "p. 26-29, 41-44")
    add_body(doc, "Structura backendului se regăsește aproape identic în cele două documente:")
    bullets = [
        "Controllerul REST primește cererea, o transformă într-un DTO și deleagă execuția; nu conține regulile principale ale afacerii.",
        "CrmFacade reprezintă punctul unic de intrare în domeniu și ascunde complexitatea serviciilor interne.",
        "Serviciile implementează regulile pentru contacte, cursuri, înscrieri, facturare, analiză și suport.",
        "Comenzile încapsulează operațiile de modificare a stării.",
        "Repository-urile oferă operații orientate spre domeniu, iar DAO-urile execută SQL JDBC direct.",
        "EventBus formează o ramură laterală, decuplată de traseul principal al cererii.",
    ]
    for item in bullets:
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "6. Șabloanele de proiectare comune", 1)
    add_source_line(doc, "p. 6-7, 10-13, 15-18", "p. 1-2, 7, 27-28, 38, 41-45, 47")
    add_patterns_table(doc)
    add_body(
        doc,
        "Lucrarea a doua nu doar enumeră aceste șabloane, ci urmărește activarea lor în același flux de cumpărare. "
        "Aceasta confirmă ideea centrală din conspect: arhitectura este construită astfel încât o operație principală să poată declanșa "
        "reacții suplimentare fără ca serviciul inițiator să le cunoască în detaliu.",
    )

    add_heading(doc, "7. Nucleul bazat pe evenimente", 1)
    add_source_line(doc, "p. 7, 10-11, 13, 15, 17-18", "p. 27-28, 38, 43-45, 47")
    add_body(
        doc,
        "Potrivirea este directă. După crearea unui contact sau a unei înscrieri, serviciul publică un eveniment. EventBus notifică "
        "observatorii înregistrați, care pot scrie în jurnalul de audit, calcula ori actualiza scorul unui lead și genera structura unei notificări. "
        "Serviciul principal rămâne independent de aceste efecte secundare.",
    )
    add_callout(
        doc,
        "Ce înseamnă simplu",
        "O singură acțiune este salvată o singură dată, după care mai multe componente pot reacționa separat. "
        "Aceasta reduce legăturile directe dintre module și face extinderea mai ușoară.",
        fill=PALE_TEAL,
        label_color=TEAL,
    )
    add_callout(
        doc,
        "Limită reală",
        "În lucrarea a doua, e-mailul și SMS-ul sunt construite prin Factory, dar metoda finală de trimitere doar înregistrează acțiunea în log. "
        "Integrarea reală cu SMTP/Twilio este marcată ca dezvoltare viitoare.",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    add_heading(doc, "8. Persistența și modelul de date", 1)
    add_source_line(doc, "p. 7, 9-12, 15-17", "p. 26-32, 42-43, 46-48")
    add_heading(doc, "Persistența", 2)
    add_body(
        doc,
        "Datele sunt stocate în MariaDB, accesate prin SQL explicit și JDBC, cu conexiuni reutilizate de HikariCP. "
        "WebSchemaInitializer rulează schema.sql la pornire. ORM-ul este evitat intenționat pentru control determinist și transparență, "
        "cu prețul unui volum mai mare de cod infrastructural.",
    )
    add_heading(doc, "Entitățile comune", 2)
    for text in [
        "Course → CourseSession → Enrollment descrie traseul educațional și înscrierea.",
        "Contact poate reprezenta o persoană individuală sau o companie.",
        "Opportunity, Activity și Employee descriu zona comercială și CRM.",
        "MediationSession este legată de Invoice pentru ședințele individuale și facturare.",
        "Evaluarea și feedbackul sunt păstrate în Enrollment, ceea ce leagă recenzia de o achiziție/înscriere reală.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_callout(
        doc,
        "Clarificare de model",
        "PDF-ul recunoaște că modelul folosește chei externe gestionate manual și denumiri moștenite, precum MediationSession. "
        "Acestea sunt limitări documentate, nu motive pentru a considera modelul inexistent.",
        fill=LIGHT_GRAY,
        label_color=DARK_BLUE,
    )

    add_heading(doc, "9. Fluxul comun de cumpărare și înscriere", 1)
    add_source_line(doc, "p. 3, 10-12, 14-15, 17-18", "p. 18, 20, 38-45")
    add_body(doc, "Acesta este cel mai bine documentat punct de corespondență. Fluxul tehnic din lucrarea a doua poate fi sistematizat astfel:")
    steps = [
        "Utilizatorul apasă butonul de înscriere din CourseCard.",
        "Frontendul trimite o cerere POST prin stratul tipizat lib/api.ts.",
        "PublicCatalogController deserializează cererea și o transmite către CrmFacade.",
        "ReviewService verifică existența și starea cursului.",
        "Contactul este căutat după e-mail; dacă lipsește, este creat înainte de Enrollment.",
        "Datele contactului trec prin lanțul de validare, iar scorul inițial este calculat prin Strategy.",
        "ContactRepository și ContactDao persistă contactul în MariaDB.",
        "Sesiunea cursului este identificată, iar EnrollmentService creează Enrollment prin Builder.",
        "EnrollmentRepository și EnrollmentDao salvează înscrierea folosind JDBC/HikariCP.",
        "EnrollmentCreatedEvent este publicat; observatorii actualizează auditul și construiesc confirmările.",
        "Răspunsul HTTP 201 revine către interfață, iar statisticile sunt transmise asincron prin SSE.",
    ]
    for step in steps:
        add_number(doc, step, number_id)
    add_callout(
        doc,
        "Idempotență confirmată",
        "Dacă înscrierea există deja, fluxul returnează înregistrarea existentă și nu creează un duplicat. "
        "Această confirmare privește în mod explicit Enrollment; afirmația din conspect despre idempotența facturii pentru ședințe nu este demonstrată la fel de clar în PDF.",
        fill=PALE_TEAL,
        label_color=TEAL,
    )

    add_heading(doc, "10. Frontendul și actualizarea în timp real", 1)
    add_source_line(doc, "p. 6, 8-9, 12", "p. 18-20, 33-37, 43, 46-47")
    for text in [
        "Next.js App Router: directoarele și page.tsx definesc rutele, iar layout.tsx furnizează cadrul comun.",
        "lib/api.ts centralizează toate cererile către backend și mapează erorile într-un contract comun.",
        "Interfețele TypeScript oglindesc DTO-urile și entitățile primite de la backend.",
        "Starea este gestionată prin mecanisme native React, Context, evenimente de fereastră și localStorage, fără Redux/MobX/Zustand.",
        "SSE folosește EventSource pentru actualizarea numărului de cursuri, cursanți și evaluări fără polling periodic.",
        "TranslationProvider și MutationObserver susțin traducerea dinamică, iar Tailwind CSS v4 gestionează tema și modul întunecat.",
    ]:
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "11. Funcționalitățile care se regăsesc", 1)
    add_source_line(doc, "p. 7-9", "p. 18-25, 34-35")
    add_heading(doc, "Portalul public / utilizator", 2)
    for text in [
        "catalog public și filtrarea cursurilor;",
        "înregistrare, autentificare și recuperarea parolei;",
        "cumpărare/înscriere și pagina My Courses;",
        "citirea, scrierea și actualizarea recenziilor;",
        "programarea și vizualizarea ședințelor cu trainerul;",
        "asistent AI, recomandări, traducere și raportarea unei probleme.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_heading(doc, "Portalul administrativ / CRM", 2)
    for text in [
        "dashboard și indicatori; contacte și leaduri;",
        "administrarea cursurilor, angajaților și înscrierilor;",
        "facturi și export PDF; pipeline de oportunități și activități;",
        "analize, rapoarte și gestionarea sesizărilor;",
        "recomandări și analize AI pentru companii și vânzări.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_callout(
        doc,
        "Limită de corespondență",
        "Detaliile foarte concrete din conspect - programul 08:00-20:00, tariful de 10 dolari/oră, reducerea exactă de 60% sau anumite ecrane administrative - "
        "nu sunt confirmate explicit în lucrarea a doua și nu trebuie prezentate ca potriviri directe.",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    add_heading(doc, "12. Integrarea AI și documentele", 1)
    add_source_line(doc, "p. 1, 8-12, 15", "p. 8, 18-22, 25-26, 32, 36, 47")
    add_heading(doc, "AI opțional", 2)
    add_body(
        doc,
        "ClaudeClient este un wrapper peste SDK-ul Anthropic. Configurarea depinde de ANTHROPIC_API_KEY. Dacă cheia lipsește, clientul rămâne neinițializat, "
        "isEnabled() întoarce false, iar apelurile produc un răspuns controlat de indisponibilitate, nu o oprire a întregii aplicații. "
        "Funcțiile comune sunt asistentul conversațional, recomandările și traducerea dinamică.",
    )
    add_heading(doc, "Importuri și exporturi", 2)
    add_body(
        doc,
        "Apache POI este folosit pentru fișiere Excel și importul angajaților, iar OpenPDF pentru facturi și exporturi PDF. "
        "Această componentă apare explicit în ambele documente.",
    )

    add_heading(doc, "13. Autentificarea și securitatea - potrivire cu precizare", 1)
    add_source_line(doc, "p. 5, 7, 9, 16-18", "p. 18-20, 32, 35, 46-48")
    add_body(
        doc,
        "Este confirmată existența rolurilor USER și ADMIN, a obiectului AuthSession, a AuthGuard și a redirecționării către portalul potrivit. "
        "Sesiunea este păstrată în localStorage, iar protecția rutelor este implementată în frontend.",
    )
    add_callout(
        doc,
        "Avertisment obligatoriu",
        "Lucrarea a doua precizează că autentificarea și autorizarea actuală sunt client-side și potrivite pentru demonstrație, nu pentru producție. "
        "Nu există încă sesiuni reale sau JWT, parole hashuite, RBAC server-side complet, MFA ori protecție industrială. "
        "Prin urmare, formularea corectă este «separare funcțională a rolurilor în interfață», nu «securitate completă».",
        fill=RED_FILL,
        label_color=RED,
    )

    add_heading(doc, "14. Ce nu trebuie preluat ca fiind confirmat", 1)
    add_body(doc, "Pentru a păstra comparația corectă, următoarele elemente din conspect nu au fost incluse în sinteza tehnică principală:")
    for text in [
        "o strategie completă de testare manuală end-to-end și o suită automată bazată pe spring-boot-starter-test; PDF-ul nu documentează un capitol de testare echivalent;",
        "Spring Validation, Lombok, SLF4J și Logback ca tehnologii folosite; nu sunt menționate explicit în lucrarea a doua;",
        "numărul aproximativ de 11.500 de linii Java și aproximativ 50 de module TypeScript/React; nu este confirmat în PDF;",
        "izolarea erorii de facturare astfel încât rezervarea să rămână salvată; nucleul evenimentelor este confirmat, dar acest comportament precis nu este demonstrat;",
        "funcționarea reală a SMTP/Twilio și Stripe/Netopia; PDF-ul le tratează ca integrări externe propuse sau încă neconectate;",
        "detaliile comerciale exacte privind programul trainerilor, tarifele și procentele de reducere.",
    ]:
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "15. Formulare tehnică consolidată", 1)
    add_callout(
        doc,
        "Variantă recomandată",
        "Textul de mai jos reunește numai ideile tehnice care sunt susținute de ambele documente și introduce precizările necesare.",
        fill=PALE_TEAL,
        label_color=TEAL,
    )
    consolidated = [
        "TrainingIT este o soluție full-stack alcătuită din două subsisteme autonome: un frontend single-page realizat cu Next.js 16, React 19, TypeScript 5 și Tailwind CSS v4 și un backend Java bazat pe Spring Boot 3.5. Cele două componente comunică exclusiv printr-un contract REST/JSON, iar actualizările publice în timp real sunt transmise prin Server-Sent Events.",
        "Backendul este organizat pe straturi. Controllerele REST sunt subțiri și deleagă cererile către CrmFacade, care reprezintă punctul unic de intrare în domeniu. Serviciile implementează logica de afaceri, comenzile încapsulează modificările de stare, repository-urile oferă operații orientate spre domeniu, iar DAO-urile execută SQL explicit prin JDBC.",
        "Persistența este realizată în MariaDB, cu un pool de conexiuni HikariCP și o schemă inițializată de aplicație. ORM-ul este evitat intenționat pentru transparență și control determinist, ceea ce implică însă mai mult cod de infrastructură și gestionarea manuală a relațiilor.",
        "Nucleul aplicației utilizează șabloane precum Facade, Singleton, Command, Repository/DAO, Builder, Strategy, Chain of Responsibility, Observer și Factory. EventBus publică evenimente precum ContactCreatedEvent și EnrollmentCreatedEvent, iar observatorii pot actualiza auditul, scorurile și notificările fără ca serviciul principal să cunoască toate efectele secundare.",
        "Fluxul de cumpărare pornește din frontend, trece prin controller, facade și servicii, identifică sau creează contactul înainte de Enrollment, validează datele, calculează scorul, persistă contactul și înscrierea, publică evenimentele și returnează răspunsul către interfață. Operația de înscriere este idempotentă: o înscriere existentă nu este duplicată.",
        "Modelul de date leagă zona educațională - Course, CourseSession și Enrollment - de zona CRM - Contact, Opportunity, Activity și Employee. Ratingul și feedbackul sunt păstrate în Enrollment, ceea ce permite asocierea recenziei cu o înscriere reală.",
        "Integrarea Claude este opțională și deservește asistentul, recomandările și traducerea. Lipsa cheii API produce degradare controlată, fără oprirea aplicației. Totuși, autentificarea actuală rămâne client-side, iar notificările și plățile externe nu sunt încă integrate complet; acestea reprezintă direcții de dezvoltare pentru o versiune de producție.",
    ]
    for paragraph in consolidated:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading(doc, "16. Concluzie finală", 1)
    add_body(
        doc,
        "Partea tehnică din primul document care se regăsește solid în al doilea document este formată din: arhitectura full-stack decuplată, "
        "stackul Next.js/Spring Boot/MariaDB, comunicarea REST și SSE, backendul stratificat, persistența JDBC fără ORM, șabloanele de proiectare, "
        "EventBus și Observer, modelul de date CRM-educațional, fluxul de cumpărare/înscriere, integrarea AI opțională și funcțiile principale ale celor două portaluri.",
    )
    add_callout(
        doc,
        "Ideea centrală",
        "TrainingIT este un singur produs funcțional construit din două aplicații tehnic independente, legate printr-un contract clar și susținute de un domeniu Java, "
        "o bază de date relațională comună și mecanisme decuplate pentru reacțiile secundare.",
        fill=GREEN_FILL,
        label_color=GREEN,
    )

    # Core properties and final save.
    props = doc.core_properties
    props.title = "Partea tehnică comună dintre conspectul TrainingIT și lucrarea de licență"
    props.subject = "Sinteză comparativă tehnică"
    props.author = "Codex"
    props.keywords = "TrainingIT, arhitectură, CRM, comparație tehnică, Spring Boot, Next.js"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
