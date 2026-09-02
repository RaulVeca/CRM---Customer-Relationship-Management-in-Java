from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from build_trainingit_comparison import (
    BLUE,
    DARK_BLUE,
    GOLD,
    GOLD_FILL,
    GREEN,
    GREEN_FILL,
    LIGHT_GRAY,
    MID_BLUE,
    MID_GRAY,
    MUTED,
    NAVY,
    PALE_BLUE,
    PALE_TEAL,
    RED,
    RED_FILL,
    TEAL,
    TEXT,
    WHITE,
    add_body,
    add_bottom_border,
    add_bullet,
    add_callout,
    add_flow_step,
    add_heading,
    add_number,
    add_page_number,
    configure_styles,
    keep_with_next,
    repeat_header,
    set_cell_shading,
    set_run_font,
    set_table_borders,
    set_table_geometry,
    setup_numbering,
    style_table_text,
)


OUTPUT = Path(r"D:\TrainingIT_site\output\documents\Partea_tehnica_TrainingIT_extrasa_din_newfile4.docx")


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        p = header.paragraphs[0]
        p.text = "TRAININGIT  |  CUM ESTE CONSTRUITĂ TEHNIC APLICAȚIA"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_run_font(run, size=8.5, bold=True, color=MUTED)
        add_bottom_border(p, color=MID_GRAY, size=4)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(p)


def add_page_reference(doc: Document, pages: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    keep_with_next(p)
    r = p.add_run("Sursa tehnică în newfile4: ")
    set_run_font(r, size=9.5, bold=True, color=MUTED)
    r = p.add_run(pages)
    set_run_font(r, size=9.5, italic=True, color=MUTED)


def add_label_body(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, color=TEXT)


def add_stack_table(doc: Document) -> None:
    rows = [
        ("Interfață", "Next.js 16, React 19, TypeScript 5, Tailwind CSS v4", "SPA autonomă, App Router, componente și interfața ambelor portaluri"),
        ("Comunicare", "HTTP + REST/JSON", "Singurul contract dintre frontend și backend"),
        ("Timp real", "Server-Sent Events (SSE)", "Actualizarea statisticilor publice fără polling și fără refresh"),
        ("Backend web", "Spring Boot 3.5", "Expune controllere REST și integrează nucleul Java"),
        ("Domeniu", "Java + design patterns", "Reguli de afaceri, servicii, comenzi, evenimente și observatori"),
        ("Persistență", "MariaDB, JDBC, HikariCP", "SQL explicit, pool de conexiuni și schemă controlată de aplicație"),
        ("AI opțional", "Anthropic Claude", "Chatbot, recomandări, asistență și traducere dinamică"),
        ("Documente", "Apache POI și OpenPDF", "Import/export Excel și generare de documente PDF"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, value in enumerate(("Zonă", "Tehnologie", "Rol în aplicație")):
        hdr[idx].text = value
        set_cell_shading(hdr[idx], PALE_BLUE)
    repeat_header(table.rows[0])
    for zone, tech, role in rows:
        cells = table.add_row().cells
        cells[0].text = zone
        cells[1].text = tech
        cells[2].text = role
    set_table_geometry(table, [1500, 3160, 4700])
    set_table_borders(table, color=MID_GRAY, size=6)
    style_table_text(table, body_size=9.0)


def add_patterns_table(doc: Document) -> None:
    rows = [
        ("Facade", "CrmFacade", "Punct unic de intrare în domeniu; ascunde serviciile și pașii interni."),
        ("Singleton", "Servicii, AppConfig, DatabaseConnection", "O singură instanță controlată pentru componentele centrale."),
        ("Command", "Operații administrative și modificări de stare", "Încapsulează operația ca unitate de execuție."),
        ("Repository + DAO", "Accesul la MariaDB", "Repository-ul vorbește în termenii domeniului; DAO-ul execută SQL/JDBC."),
        ("Builder", "Construirea Enrollment", "Asamblează obiectul complex și îl inițializează cu status CONFIRMED."),
        ("Observer + EventBus", "ContactCreatedEvent, EnrollmentCreatedEvent", "Declanșează reacții secundare fără a cupla serviciul principal de observatori."),
        ("Factory", "Notificări e-mail și SMS", "Construiește mesaje adaptate canalului prin aceeași interfață."),
        ("Strategy", "Scorarea leadurilor B2C/B2B", "Selectează algoritmul potrivit tipului de contact."),
        ("Chain of Responsibility", "Validarea contactului", "Parcurge succesiv e-mail, telefon, tip de contact și GDPR și acumulează erorile."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, value in enumerate(("Șablon", "Unde apare", "Ce rezolvă")):
        table.rows[0].cells[idx].text = value
        set_cell_shading(table.rows[0].cells[idx], PALE_BLUE)
    repeat_header(table.rows[0])
    for pattern, where, role in rows:
        cells = table.add_row().cells
        cells[0].text = pattern
        cells[1].text = where
        cells[2].text = role
    set_table_geometry(table, [1880, 3050, 4430])
    set_table_borders(table, color=MID_GRAY, size=6)
    style_table_text(table, body_size=8.8)


def add_source_map(doc: Document) -> None:
    rows = [
        ("26-28", "Arhitectura generală și arhitectura backend"),
        ("29-31", "Entități, subdomenii, relații și enumerări"),
        ("32", "Persistență, AI, raportare și securitate"),
        ("32-37", "Arhitectura frontend"),
        ("38-45", "Studiul de caz al cumpărării și înscrierii"),
        ("46-48", "Contribuții, limite și dezvoltări necesare"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Pagini"
    table.rows[0].cells[1].text = "Conținut tehnic relevant"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, PALE_BLUE)
    repeat_header(table.rows[0])
    for pages, content in rows:
        cells = table.add_row().cells
        cells[0].text = pages
        cells[1].text = content
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, [1500, 7860])
    set_table_borders(table, color=MID_GRAY, size=6)
    style_table_text(table, body_size=9.3)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_id, number_id = setup_numbering(doc)
    _, recap_number_id = setup_numbering(doc)

    # Editorial cover.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("\n\n\n\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("EXTRAS TEHNIC DIN LUCRAREA DE LICENȚĂ")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Cum este construită tehnic")
    set_run_font(r, size=29, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("aplicația TrainingIT")
    set_run_font(r, size=29, bold=True, color=NAVY)
    add_bottom_border(p, color=BLUE, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    r = p.add_run("Explicație tehnică structurată după paginile 26-48 din newfile4")
    set_run_font(r, size=13.5, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Arhitectură • straturi • date • flux de înscriere • design patterns • limite reale")
    set_run_font(r, size=10.5, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(54)
    r = p.add_run("Material pentru prezentarea la susținerea lucrării de licență")
    set_run_font(r, size=11, bold=True, color=NAVY)

    doc.add_page_break()

    add_heading(doc, "1. Ideea tehnică centrală", 1)
    add_page_reference(doc, "p. 26-28 și p. 46-47")
    add_callout(
        doc,
        "De reținut",
        "TrainingIT este un singur produs funcțional, dar este construit din două aplicații autonome: un frontend Next.js și un backend Java/Spring Boot. Ele nu împart cod, sesiune sau componente interne; comunică exclusiv prin HTTP și mesaje JSON.",
        fill=GREEN_FILL,
        label_color=GREEN,
    )
    add_body(
        doc,
        "Frontendul gestionează interfața, navigarea și starea locală a utilizatorului. Backendul implementează regulile de afaceri, accesul la date și reacțiile administrative. MariaDB păstrează atât informațiile educaționale, cât și evidența CRM. Această separare permite modificarea interfeței fără rescrierea domeniului Java și modificarea backendului fără recompilarea frontendului, atât timp cât este păstrat contractul REST/JSON.",
    )

    add_heading(doc, "Traseul unei cereri", 2)
    for label, detail, fill in [
        ("1. Frontend", "Next.js / React colectează acțiunea utilizatorului", PALE_BLUE),
        ("2. REST/JSON", "lib/api.ts trimite cererea HTTP către backend", LIGHT_GRAY),
        ("3. Controller", "deserializare, validare web și delegare; fără logică de afaceri", PALE_BLUE),
        ("4. CrmFacade", "punct unic de intrare și coordonare a domeniului", PALE_TEAL),
        ("5. Servicii / comenzi", "aplică regulile aplicației și modifică starea", LIGHT_GRAY),
        ("6. Repository / DAO", "transformă operația de domeniu în SQL explicit", PALE_BLUE),
        ("7. JDBC / HikariCP / MariaDB", "execută și persistă datele", PALE_TEAL),
        ("8. EventBus", "ramifică reacțiile secundare: audit, notificări și statistici", GOLD_FILL),
    ]:
        add_flow_step(doc, label, detail, fill)

    add_heading(doc, "2. Tehnologiile care alcătuiesc sistemul", 1)
    add_page_reference(doc, "p. 26, p. 32-37 și p. 46")
    add_stack_table(doc)
    add_callout(
        doc,
        "Alegere arhitecturală",
        "Spring Boot este folosit în principal ca strat web. Nucleul de domeniu Java își gestionează singur configurația, conexiunile, schema și tranzacțiile, ceea ce reduce dependența logicii de afaceri față de framework.",
        fill=PALE_TEAL,
        label_color=TEAL,
    )

    add_heading(doc, "3. Cum este construit backendul", 1)
    add_page_reference(doc, "p. 26-32 și p. 41-45")
    add_label_body(doc, "Stratul web", "Aproximativ șaisprezece controllere REST din crm.web.controller transformă cererile HTTP în apeluri de domeniu și returnează DTO-uri. Ele sunt intenționat subțiri.")
    add_label_body(doc, "Facade", "Toate operațiile importante intră prin CrmFacade. Controllerul nu trebuie să cunoască serviciile, comenzile sau repository-urile care vor fi activate.")
    add_label_body(doc, "Serviciile", "Componentele crm.service.* conțin regulile pentru contacte, cursuri, înscrieri, facturare, analiză, recenzii și suport. Serviciile centrale sunt configurate ca Singleton.")
    add_label_body(doc, "Comenzile", "Operațiile care modifică starea pot fi încapsulate prin Command și executate prin CommandInvoker, păstrând clară responsabilitatea fiecărei modificări.")
    add_label_body(doc, "Persistența", "Repository-urile oferă operații orientate spre domeniu; DAO-urile execută PreparedStatement-uri și SQL nativ prin JDBC. Nu sunt folosite JPA sau Hibernate.")
    add_label_body(doc, "Evenimentele", "După o modificare importantă, serviciul publică un eveniment. EventBus notifică observatori independenți, astfel încât operația principală nu cunoaște toate reacțiile secundare.")

    add_heading(doc, "Contractul unitar de eroare", 2)
    add_body(doc, "GlobalExceptionHandler, implementat cu @RestControllerAdvice, transformă excepțiile domeniului într-un răspuns ApiError cu status, mesaj, cale și detalii:")
    for item in [
        "resursă inexistentă → HTTP 404 Not Found;",
        "validare eșuată → HTTP 400 Bad Request;",
        "regulă de afaceri încălcată → HTTP 422 Unprocessable Entity;",
        "eroare de persistență → HTTP 500 Internal Server Error;",
        "serviciul AI indisponibil → HTTP 503 Service Unavailable.",
    ]:
        add_bullet(doc, item, bullet_id)
    add_callout(doc, "Formulare orală", "Controllerul primește cererea, CrmFacade găsește traseul, serviciul aplică regula, iar repository-ul și DAO-ul salvează rezultatul.", fill=GREEN_FILL, label_color=GREEN)

    add_heading(doc, "4. Cum este construit frontendul", 1)
    add_page_reference(doc, "p. 32-37")
    add_label_body(doc, "Aplicație autonomă", "Frontendul se află într-un director separat, rulează independent și tratează backendul ca pe un serviciu extern accesibil numai prin REST/JSON.")
    add_label_body(doc, "App Router", "Fiecare director din src/app definește o rută. Sunt separate paginile publice, autentificarea, zonele utilizatorului și ramura administrativă /admin/*.")
    add_label_body(doc, "Root layout", "app/layout.tsx furnizează cadrul comun: header, navigație, selector de limbă, temă, footer, chatbot, raportarea problemelor, TranslationProvider și AuthGuard.")
    add_label_body(doc, "Gateway API", "lib/api.ts centralizează get/post/put/patch/delete/upload, adresa backendului, Content-Type, cache: no-store și conversia erorilor în ApiError.")
    add_label_body(doc, "Contract tipizat", "lib/types.ts oglindește DTO-urile serverului prin interfețe TypeScript precum PublicCourse, PublicStats, MyPurchase și AuthSession.")
    add_label_body(doc, "Starea aplicației", "useState și useEffect gestionează starea locală; localStorage păstrează sesiunea, limba și tema; evenimentele globale și React Context propagă schimbările fără Redux, MobX sau Zustand.")
    add_label_body(doc, "SSE", "EventSource se conectează la /api/public/stats/stream și actualizează automat cursurile, cursanții, ratingul mediu și recenziile.")

    add_heading(doc, "5. Persistența și modelul de date", 1)
    add_page_reference(doc, "p. 29-32 și p. 42-43")
    add_body(doc, "DatabaseConnection este un Singleton care gestionează poolul HikariCP. La pornire, WebSchemaInitializer.ensureTables() execută schema.sql, iar nucleul Java își creează și își controlează schema în MariaDB. Spring DataSource este dezactivat intenționat.")
    add_callout(doc, "De ce fără ORM", "Legăturile sunt păstrate ca identificatori Long și SQL-ul rămâne explicit. Avantajul este controlul determinist și transparența; costul este mai mult cod și gestionarea manuală a relațiilor.", fill=PALE_TEAL, label_color=TEAL)

    add_heading(doc, "Cele trei subdomenii", 2)
    for item in [
        "Educațional: Course → CourseSession → Enrollment. Ratingul și feedbackul sunt păstrate în Enrollment.",
        "CRM și vânzări: Contact, Opportunity, Activity și Employee.",
        "Ședințe individuale și facturare: MeditationSession legată de Invoice.",
    ]:
        add_bullet(doc, item, bullet_id)
    add_body(doc, "Contact poate reprezenta o persoană individuală sau o companie. Clientul/cursantul nu are o entitate Learner separată, ci este modelat prin Contact. Stările și categoriile sunt controlate prin enumerări pentru rol, tip de contact, lead, oportunitate, înscriere, plată, categorie de curs, mod de livrare și stare de sesiune.")

    add_heading(doc, "6. Cum sunt realizate funcțiile prezentate în conspect", 1)
    add_page_reference(doc, "p. 29-37 și p. 38-45")
    mappings = [
        ("Catalog și cumpărare", "CourseCard.tsx afișează cursul și trimite POST către /api/public/courses/{id}/purchase prin lib/api.ts."),
        ("Cursurile cumpărate", "Înscrierea este persistată ca Enrollment asociat unui Contact și unei CourseSession."),
        ("Recenzii", "rating și feedback sunt câmpuri ale Enrollment, ceea ce leagă evaluarea de înscriere."),
        ("CRM administrativ", "ramura /admin/* oferă ecrane pentru contacte, cursuri, achiziții, facturare, analize, angajați și sesizări; backendul le servește prin același domeniu."),
        ("Ședințe individuale", "entitățile MeditationSession și Invoice modelează rezervarea cu trainerul și factura asociată."),
        ("Raportare", "ReportService și ReportController generează Excel/PDF, iar EmployeeExcelImporter realizează importul în masă al angajaților."),
        ("Asistent și recomandări", "ClaudeClient și AiService oferă chatbot, recomandări de curs, sprijin comercial și traducere."),
        ("Actualizări live", "PublicStatsBroadcaster recalculează statisticile și le transmite prin SSE numai când există un client conectat."),
    ]
    for label, text in mappings:
        add_label_body(doc, label, text)

    add_heading(doc, "7. Fluxul tehnic al cumpărării și înscrierii", 1)
    add_page_reference(doc, "p. 38-45")
    add_callout(doc, "De ce este fluxul reprezentativ", "O singură apăsare a butonului Enroll traversează aproape toate straturile aplicației și activează majoritatea șabloanelor de proiectare.", fill=GREEN_FILL, label_color=GREEN)
    purchase_steps = [
        "CourseCard.tsx verifică dacă utilizatorul este deja înscris și afișează Enroll sau Enrolled.",
        "La click, frontendul trimite prin lib/api.ts un POST cu e-mailul, prenumele și numele utilizatorului.",
        "PublicCatalogController.purchase() deserializează cererea și o deleagă imediat către CrmFacade.",
        "CrmFacade transmite operația către ReviewService.purchaseCourse().",
        "ReviewService verifică existența și starea activă a cursului.",
        "findOrCreateContact() caută contactul după e-mail; dacă lipsește, îl creează în CRM.",
        "Contactul nou trece prin lanțul Email → Phone → ContactType → GDPR; Strategy calculează scorul inițial, iar Repository/DAO îl persistă.",
        "ContactCreatedEvent este publicat, iar WelcomeEmailObserver construiește notificarea de bun venit.",
        "Serviciul verifică idempotența; dacă există deja Enrollment pentru contact și curs, returnează înregistrarea existentă.",
        "resolveSession(course) selectează sesiunea, iar EnrollmentService validează contactul și sesiunea.",
        "Enrollment.builder() creează înscrierea cu status CONFIRMED; EnrollmentRepository și EnrollmentDao execută INSERT în MariaDB.",
        "Leadul contactului devine ENROLLED, apoi EnrollmentCreatedEvent declanșează auditul și confirmarea.",
        "Răspunsul urcă prin serviciu, facade și controller; frontendul primește HTTP 201 și schimbă butonul în Enrolled.",
        "PublicStatsBroadcaster detectează modificarea și transmite valorile noi prin SSE către paginile deschise.",
    ]
    for step in purchase_steps:
        add_number(doc, step, number_id)

    add_callout(doc, "Idempotență", "Repetarea aceleiași cereri nu creează o a doua înscriere; este returnată înregistrarea existentă.", fill=PALE_TEAL, label_color=TEAL)

    add_heading(doc, "8. Șabloanele de proiectare care susțin fluxul", 1)
    add_page_reference(doc, "p. 27-28 și p. 38-45")
    add_patterns_table(doc)
    add_body(doc, "EventBus formează o ramură laterală față de traseul principal. Serviciul publică faptul că s-a produs o schimbare, iar observatorii decid independent ce reacții execută. Astfel, auditul, notificarea și recalcularea statisticilor pot fi adăugate sau modificate fără rescrierea operației de înscriere.")

    add_heading(doc, "9. AI, traducere, temă și notificări", 1)
    add_page_reference(doc, "p. 32, p. 36 și p. 41-43")
    add_label_body(doc, "AI cu degradare controlată", "ClaudeClient este un wrapper peste SDK-ul Anthropic. Dacă ANTHROPIC_API_KEY lipsește, isEnabled() întoarce false și este produsă o eroare controlată 503, fără oprirea aplicației.")
    add_label_body(doc, "Traducere dinamică", "TranslationProvider gestionează limba. Motorul extrage textul englez, îl trimite în loturi de maximum 25 de șiruri, limitează concurența la trei apeluri și reaplică traducerile prin MutationObserver. Rezultatele sunt memorate în localStorage.")
    add_label_body(doc, "Temă", "Tailwind CSS v4 definește tokenurile vizuale. Un script rulat înainte de hidratarea React aplică imediat tema memorată sau preferința sistemului, evitând afișarea temporară a temei greșite.")
    add_label_body(doc, "Notificări", "NotificationService folosește fabrici separate pentru e-mail și SMS. E-mailul este canalul principal, iar SMS-ul este construit numai dacă există număr de telefon.")
    add_callout(doc, "Limită actuală", "Metoda finală send() nu transmite încă mesaje reale; scrie doar în log. Integrarea SMTP/Twilio este marcată ca dezvoltare viitoare.", fill=GOLD_FILL, label_color=GOLD)

    add_heading(doc, "10. Ce este implementat demonstrativ, nu pentru producție", 1)
    add_page_reference(doc, "p. 32 și p. 46-48")
    for item in [
        "Autentificarea și rolul sunt păstrate în localStorage și controlate în principal de AuthGuard și de layoutul administrativ.",
        "Nu există încă sesiuni reale sau JWT, parole stocate prin hashing, RBAC complet pe server ori autentificare multifactor.",
        "CORS este restricționat la originea serverului Next.js, dar aceasta nu înlocuiește securitatea server-side.",
        "Fluxul de cumpărare confirmă instant Enrollment și nu conține un checkout real; Stripe sau Netopia sunt propuneri viitoare.",
        "Notificările e-mail/SMS sunt construite corect ca structură, dar nu sunt conectate la furnizori externi.",
        "GdprValidator conține o condiție inertă, indicată în lucrare drept un silent bug ce trebuie corectat.",
        "Numele MeditationSession este moștenit și ar trebui schimbat în TrainerSession sau Booking.",
        "Executarea directă a schema.sql ar trebui înlocuită cu migrații controlate, precum Flyway sau Liquibase, iar operațiile pe mai multe tabele necesită atomicitate mai strictă.",
    ]:
        add_bullet(doc, item, bullet_id)
    add_callout(doc, "Formulare corectă la susținere", "Aplicația demonstrează o arhitectură completă end-to-end și design patterns aplicate real, dar securitatea, plățile și mesajele externe trebuie consolidate pentru producție.", fill=RED_FILL, label_color=RED)

    doc.add_page_break()
    add_heading(doc, "11. Explicația orală, în aproximativ 90 de secunde", 1)
    add_body(
        doc,
        "TrainingIT este construit din două aplicații autonome. Frontendul este o aplicație Next.js și React, iar backendul este un nucleu Java expus prin Spring Boot. Cele două comunică numai prin REST și JSON. În backend, controllerul este subțire și deleagă cererea către CrmFacade. De acolo, serviciile aplică regulile de afaceri, iar repository-urile și DAO-urile persistă datele în MariaDB prin JDBC și HikariCP. Nu este folosit un ORM, pentru ca SQL-ul și traseul datelor să rămână explicite.",
    )
    add_body(
        doc,
        "Partea distinctivă este nucleul bazat pe evenimente. După o acțiune importantă, cum este înscrierea la un curs, EventBus notifică observatori independenți care pot scrie în audit, pregăti confirmări și actualiza statisticile. În fluxul de cumpărare sunt folosite împreună Facade, Singleton, Repository/DAO, Builder, Strategy, Chain of Responsibility, Observer și Factory. Frontendul primește răspunsul HTTP, iar statisticile sunt actualizate în timp real prin SSE.",
    )
    add_body(
        doc,
        "AI-ul Claude este opțional și deservește chatbotul, recomandările și traducerea; dacă lipsește cheia, restul aplicației continuă să funcționeze. Implementarea este completă ca demonstrație academică, însă autentificarea este încă client-side, plata este simulată, iar mesajele e-mail și SMS sunt doar înregistrate în log.",
    )

    add_heading(doc, "Cele 7 idei pe care trebuie să le reții", 2)
    for item in [
        "Două aplicații autonome, un singur produs: Next.js în față, Java/Spring Boot în spate.",
        "Comunicare exclusiv prin REST/JSON; actualizări live prin SSE.",
        "Controller → CrmFacade → servicii → repository → DAO/JDBC → MariaDB.",
        "Fără ORM: SQL explicit și relații gestionate manual.",
        "EventBus și Observer transformă o singură acțiune în mai multe reacții independente.",
        "Fluxul de înscriere este idempotent și activează mai multe design patterns.",
        "AI-ul este opțional; securitatea, plata și trimiterea notificărilor trebuie maturizate pentru producție.",
    ]:
        add_number(doc, item, recap_number_id)

    add_heading(doc, "12. Unde se găsește această parte în lucrarea mare", 1)
    add_source_map(doc)
    add_callout(doc, "Concluzie", "Pentru prezentarea tehnică a produsului din conspect, nucleul relevant al lucrării newfile4 este concentrat în paginile 26-48, cu studiul de caz esențial al înscrierii în paginile 38-45.", fill=GREEN_FILL, label_color=GREEN)

    props = doc.core_properties
    props.title = "Cum este construită tehnic aplicația TrainingIT"
    props.subject = "Extras tehnic din newfile4 pentru susținerea lucrării de licență"
    props.author = "Codex"
    props.keywords = "TrainingIT, arhitectură, Spring Boot, Next.js, MariaDB, design patterns"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
