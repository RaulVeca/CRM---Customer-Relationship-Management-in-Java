from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\TrainingIT_site\output\documents\Conspect_TrainingIT.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 38, 48)
MUTED = RGBColor(95, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def set_run_font(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph(style="Normal")
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.insert(0, num_pr)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_key_point(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + ": ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_term_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0]
    header.cells[0].text = "Termen"
    header.cells[1].text = "Explicație în contextul TrainingIT"
    for cell in header.cells:
        shade_cell(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, color=DARK_BLUE)
    set_repeat_table_header(header)
    for term, explanation in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        cells[0].text = term
        cells[1].text = explanation
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True, color=INK)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, color=INK)
    set_table_geometry(table, [2700, 6660])
    return table


def chapter(doc, number, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = False
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Capitolul {number}. {title}")
    set_run_font(r, size=16, color=BLUE, bold=True)


def heading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(text)
    set_run_font(r, size=13, color=BLUE, bold=True)


def subheading(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    set_run_font(r, size=12, color=DARK_BLUE, bold=True)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        s = styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        s = styles[list_name]
        s.font.name = "Calibri"
        s.font.size = Pt(11)
        s.paragraph_format.left_indent = Inches(0.375)
        s.paragraph_format.first_line_indent = Inches(-0.188)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("TRAININGIT  |  CONSPECT")
    set_run_font(hr, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    prefix = fp.add_run("Pagina ")
    set_run_font(prefix, size=9, color=MUTED)
    add_page_field(fp)
    for run in fp.runs:
        set_run_font(run, size=9, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)

    # Editorial cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONSPECT")
    set_run_font(r, size=12, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("TrainingIT")
    set_run_font(r, size=31, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Sinteza în limba română a lucrării de licență")
    set_run_font(r, size=15, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(90)
    r = p.add_run("Aplicație web pentru comercializarea cursurilor IT și administrarea relațiilor cu clienții")
    set_run_font(r, size=11, color=INK, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document de recapitulare • 2026")
    set_run_font(r, size=10, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)

    # Contents
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run("Cuprins orientativ")
    set_run_font(r, size=16, color=BLUE, bold=True)
    add_numbered(doc, [
        "Introducere și obiective",
        "Problema abordată și soluțiile similare",
        "Arhitectura sistemului",
        "Portalurile și funcționalitățile aplicației",
        "Tehnologiile utilizate",
        "Implementarea și modelul de date",
        "Testarea aplicației",
        "Scenarii de utilizare",
        "Concluzii și dezvoltări viitoare",
        "Glosar de termeni",
    ])
    add_key_point(
        doc,
        "Notă",
        "Documentul reunește conspectele realizate în conversație și păstrează accentul pe ideile esențiale ale aplicației TrainingIT.",
    )

    chapter(doc, 1, "Introducere și obiective")
    heading(doc, "Contextul domeniului")
    add_body(doc, "Sectorul IT este unul dintre domeniile cu cea mai rapidă dezvoltare, iar învățarea continuă a devenit o necesitate profesională. Cererea pentru cursuri online de programare a crescut, iar furnizorii de formare își comercializează tot mai frecvent serviciile prin platforme web.")
    add_body(doc, "O companie de training online trebuie să răspundă simultan la două nevoi: să ofere un magazin online atractiv și să gestioneze relația comercială cu persoanele care folosesc platforma.")
    add_bullets(doc, [
        "Platforma e-commerce permite descoperirea cursurilor, înscrierea, plata și solicitarea de asistență.",
        "Sistemul CRM urmărește clienții, etapele procesului de cumpărare, achizițiile, veniturile și problemele raportate.",
        "În soluțiile tradiționale, cele două sisteme sunt separate și conectate prin integrări fragile.",
    ])

    heading(doc, "Soluția TrainingIT")
    add_body(doc, "TrainingIT integrează într-un singur produs o platformă de comercializare a cursurilor și un sistem CRM. Pentru utilizatori, aplicația funcționează ca o piață online modernă; pentru echipa companiei, aceeași aplicație oferă un portal administrativ complet.")
    add_bullets(doc, [
        "Clientul își poate crea un cont, consulta catalogul, cumpăra cursuri, publica recenzii și programa ședințe individuale.",
        "Administratorul gestionează contacte, clienți potențiali, vânzări, facturi, analize, angajați și solicitări de asistență.",
        "Acțiunile clientului declanșează automat procese precum emiterea facturii, recalcularea scorului și înregistrarea în jurnalul de audit.",
    ])
    add_key_point(doc, "Ideea centrală", "O singură acțiune efectuată de client poate produce mai multe reacții automate în CRM, fără introducerea manuală repetată a datelor.")

    heading(doc, "Obiectivul principal")
    add_body(doc, "Obiectivul lucrării este proiectarea și implementarea unei aplicații web apropiate de cerințele unui produs real, destinată comercializării cursurilor de programare și înregistrării automate a activității clienților într-un sistem CRM.")
    add_bullets(doc, [
        "realizarea unui portal public pentru cursuri, achiziții, recenzii și programări;",
        "dezvoltarea unui portal administrativ pentru activitățile operaționale și CRM;",
        "organizarea logicii serverului prin șabloane de proiectare și evenimente;",
        "integrarea opțională a unui asistent AI și a traducerii dinamice;",
        "stocarea persistentă a datelor într-o bază de date relațională.",
    ])

    chapter(doc, 2, "Problema abordată și soluțiile similare")
    heading(doc, "Așteptările cursantului")
    add_body(doc, "Un potențial cursant se așteaptă la o experiență asemănătoare unui magazin online modern. El trebuie să poată identifica ușor cursurile, prețurile, nivelul de dificultate și opiniile altor participanți.")
    add_bullets(doc, [
        "crearea rapidă a unui cont și cumpărarea simplă a unui curs;",
        "programarea unei ședințe individuale într-un interval convenabil;",
        "raportarea ușoară a problemelor;",
        "navigare clară, calendare corecte și recenzii credibile.",
    ])

    heading(doc, "Nevoile operaționale ale companiei")
    add_body(doc, "Compania trebuie să identifice potențialii clienți, să le urmărească apropierea de cumpărare și să concentreze eforturile de vânzare asupra oportunităților relevante.")
    add_bullets(doc, [
        "păstrarea istoricului achizițiilor și generarea facturilor corecte;",
        "aplicarea reducerilor negociate pentru angajații companiilor partenere;",
        "exportarea datelor pentru contabilitate;",
        "monitorizarea clienților pierduți, a înscrierilor abandonate și a cursurilor slab accesate;",
        "gestionarea clienților corporativi și a echipelor lor de angajați.",
    ])

    heading(doc, "Problema integrării")
    add_body(doc, "Platforma e-commerce și CRM-ul sunt, de regulă, sisteme separate. Fiecare punct de sincronizare poate produce neconcordanțe: achiziții absente din CRM, reduceri înregistrate doar pe factură sau solicitări de asistență care nu ajung la echipa comercială.")
    add_body(doc, "Verificarea și corectarea manuală a acestor diferențe este lentă și predispusă la erori. TrainingIT elimină această reconciliere prin utilizarea unui singur sistem și a unui model comun de date.")

    heading(doc, "Abordări similare")
    subheading(doc, "Platforme generale de cursuri")
    add_body(doc, "Udemy și Coursera oferă cataloage, recenzii și procese de cumpărare bine dezvoltate. Totuși, sunt platforme agregatoare, iar furnizorul nu controlează integral relația cu clientul și nu dispune de propriul CRM integrat.")
    subheading(doc, "Sisteme de management al învățării")
    add_body(doc, "Moodle este orientat către livrarea și monitorizarea conținutului educațional, precum înscrieri, note și progres. Funcțiile de comerț electronic și CRM sunt secundare și sunt adăugate, de regulă, prin extensii.")
    subheading(doc, "Platforme CRM")
    add_body(doc, "Salesforce și HubSpot oferă funcții performante pentru contacte, etape de vânzare, scoruri și rapoarte, dar nu includ în mod implicit catalog de cursuri, calendare ale trainerilor sau un magazin public.")
    subheading(doc, "Platforme e-commerce găzduite local")
    add_body(doc, "WooCommerce poate comercializa cursuri, dar le tratează ca produse obișnuite și nu gestionează etapele clienților potențiali, ședințele individuale sau instruirea corporativă.")
    add_key_point(doc, "Diferențiator", "TrainingIT combină magazinul și CRM-ul într-o singură aplicație, adaptând funcțiile comerciale la domeniul formării profesionale.")

    heading(doc, "Separarea rolurilor")
    add_body(doc, "TrainingIT este o aplicație unitară cu două portaluri specifice rolurilor, care folosesc același backend Spring Boot și aceeași bază de date MariaDB.")
    add_bullets(doc, [
        "USER: cursant individual sau angajat al unei companii, cu acces la portalul clientului;",
        "ADMIN: membru al echipei TrainingIT, cu acces la portalul administrativ CRM;",
        "încercarea de a accesa o rută rezervată celuilalt rol produce redirecționarea către portalul propriu.",
    ])

    chapter(doc, 5, "Arhitectura sistemului")
    add_body(doc, "TrainingIT este organizată în trei niveluri principale, care comunică prin API-uri REST și printr-un flux de evenimente în timp real. Sistemul utilizează o bază de date relațională și poate comunica opțional cu un serviciu extern de inteligență artificială.")

    heading(doc, "Nivelul client")
    add_body(doc, "Interfața rulează în browser și este realizată cu Next.js și React. Rolul utilizatorului este păstrat în sesiunea locală, iar mecanismele de protejare a rutelor controlează accesul la pagini.")
    add_bullets(doc, [
        "comunicare cu serverul printr-un client REST tipizat;",
        "actualizarea statisticilor prin Server-Sent Events;",
        "gestionarea temei luminoase sau întunecate;",
        "traducerea dinamică a interfeței.",
    ])

    heading(doc, "Nivelul backend")
    add_body(doc, "Backendul este un serviciu Spring Boot care expune endpointuri REST sub ruta /api/. Controllerele validează datele primite și deleagă operațiile către nivelul de domeniu.")
    add_bullets(doc, [
        "CrmFacade oferă un punct central pentru operațiile CRM principale;",
        "servicii specializate gestionează autentificarea, AI-ul, trainerii și rapoartele;",
        "un broadcaster dedicat recalculează și transmite statisticile publice prin SSE.",
    ])

    heading(doc, "Nivelul de domeniu")
    add_body(doc, "Nivelul de domeniu reprezintă componenta centrală a sistemului și demonstrează utilizarea șabloanelor de proiectare orientate pe obiecte.")
    add_bullets(doc, [
        "operațiile de scriere sunt încapsulate sub formă de comenzi;",
        "regulile aplicației sunt implementate în servicii de tip Singleton;",
        "EventBus notifică observatorii responsabili de audit, e-mailuri, facturi și scoruri;",
        "repository-urile și obiectele DAO gestionează persistența datelor.",
    ])

    heading(doc, "Nivelul de persistență")
    add_body(doc, "Datele sunt stocate într-o bază de date MariaDB, accesată prin MariaDB JDBC și un pool de conexiuni HikariCP. Schema include atât tabele CRM, cât și tabelele funcționalităților web și este creată sau actualizată la pornirea aplicației.")
    add_key_point(doc, "Persistență", "Datele rămân salvate după oprirea și repornirea aplicației.")

    heading(doc, "Controlul accesului")
    add_body(doc, "Tipul contului este stabilit la autentificare prin verificarea adresei de e-mail în evidențele administratorilor, contactelor și angajaților. Aplicația creează o sesiune USER sau ADMIN și direcționează utilizatorul către portalul corespunzător.")
    add_body(doc, "Pentru un angajat, sistemul identifică sau creează un contact în portal și aplică automat reducerea asociată companiei partenere.")

    chapter(doc, 6, "Portalurile și funcționalitățile aplicației")
    heading(doc, "Autentificare și înregistrare")
    add_body(doc, "Utilizatorii neautentificați pot accesa ecranele Log in, Register și Forgot password. Formularul de înregistrare este împărțit în secțiuni pentru date personale, contact, adresă, profil educațional, parolă și consimțământ GDPR.")
    add_body(doc, "Înregistrarea creează în CRM un contact INDIVIDUAL, calculează scorul de potențial client și autentifică automat utilizatorul. La conectare, aplicația verifică e-mailul și parola, stabilește rolul și alege portalul de destinație.")

    heading(doc, "Portalul clientului")
    subheading(doc, "Catalogul de cursuri")
    add_body(doc, "Catalogul afișează un mesaj personalizat și indicatori actualizați în timp real: numărul cursurilor, numărul cursanților, evaluarea medie și disponibilitatea mentoratului individual.")
    add_bullets(doc, [
        "un chestionar AI opțional recomandă cursuri după interese, nivel și obiectiv;",
        "cursurile sunt afișate într-o grilă filtrabilă;",
        "cardurile prezintă descrierea, evaluarea medie și numărul recenziilor;",
        "butonul Enroll permite cumpărarea rapidă;",
        "prima extindere a cardului este înregistrată pentru analiza accesărilor.",
    ])

    subheading(doc, "Cursurile utilizatorului și recenziile")
    add_body(doc, "Pagina My Courses reprezintă biblioteca personală a cursantului și este vizibilă numai proprietarului. Pentru fiecare curs cumpărat poate fi publicată o singură recenzie, formată dintr-o evaluare între 1 și 5 și un comentariu.")
    add_body(doc, "Tot din această pagină poate fi început procesul de programare a unei ședințe individuale.")

    subheading(doc, "Programarea ședințelor individuale")
    add_numbered(doc, [
        "Selectarea trainerului.",
        "Alegerea zilei.",
        "Stabilirea orei de început și a duratei.",
        "Confirmarea și efectuarea plății.",
    ])
    add_body(doc, "Trainerii sunt disponibili de luni până sâmbătă, între 08:00 și 20:00. Zilele trecute, duminicile și zilele complet ocupate sunt dezactivate. Tariful este de 10 dolari pe oră, iar angajații companiilor partenere primesc automat reducerea configurată.")
    add_bullets(doc, [
        "confirmarea rezervă ședința;",
        "factura este generată automat;",
        "intervalul devine indisponibil;",
        "anularea din My Sessions eliberează intervalul.",
    ])

    subheading(doc, "Butoanele flotante")
    add_bullets(doc, [
        "butonul din colțul din dreapta jos deschide asistentul conversațional Claude;",
        "butonul cu steag deschide formularul Report an issue;",
        "sesizarea ajunge imediat în secțiunea Issues și conține numele și e-mailul utilizatorului.",
    ])

    heading(doc, "Portalul de administrare")
    add_body(doc, "Portalul este disponibil numai unei sesiuni ADMIN și este organizat în șapte secțiuni.")
    add_bullets(doc, [
        "Contacts: persoane și companii, etape de lead, scor automat, căutare și export;",
        "Courses: prezentarea internă, tabelară, a catalogului;",
        "Purchases: istoricul înscrierilor, cursantul, cursul, data și evaluarea;",
        "Invoices: facturile generate pentru ședințe, inclusiv reducerile angajaților;",
        "Analytics: indicatori de sănătate comercială, interpretări și grafice;",
        "Employees: angajații companiilor, import Excel și recomandări AI pentru echipă;",
        "Issues: sesizările clienților, care pot fi rezolvate, redeschise sau șterse cu confirmare.",
    ])

    heading(doc, "Funcționalități disponibile pe toate paginile")
    add_bullets(doc, [
        "traducerea în timp real a conținutului, inclusiv a datelor încărcate dinamic, cu păstrarea numelor, adreselor de e-mail și numerelor;",
        "schimbarea instantanee a temei luminoase sau întunecate și memorarea locală a preferinței.",
    ])

    heading(doc, "Funcționalități AI")
    add_bullets(doc, [
        "recomandări pentru angajații companiilor partenere;",
        "recomandări individuale pe baza intereselor, experienței și obiectivelor;",
        "chatbot care folosește catalogul drept context;",
        "traducerea dinamică a interfeței.",
    ])
    add_body(doc, "Integrarea include instrucțiunile transmise modelului, endpointurile REST, procesarea răspunsurilor JSON și interfețele aferente. Dacă AI-ul nu este configurat, funcțiile sale dispar fără a afecta restul aplicației.")

    chapter(doc, 7, "Tehnologiile utilizate")
    heading(doc, "Front-end")
    add_bullets(doc, [
        "Next.js 16 și React 19: componente reutilizabile pentru ambele portaluri;",
        "TypeScript 5: tipizare statică și identificarea mai rapidă a erorilor;",
        "Tailwind CSS 4: tema vizuală personalizată și modul întunecat.",
    ])

    heading(doc, "Back-end")
    add_bullets(doc, [
        "Java 17: implementarea nivelului web și a domeniului;",
        "Spring Boot 3.5: configurarea componentelor, controllere REST și SSE;",
        "Spring Validation: validarea declarativă a cererilor;",
        "Lombok: reducerea codului repetitiv din entități și DTO-uri;",
        "SLF4J și Logback: înregistrarea mesajelor, avertismentelor și erorilor.",
    ])

    heading(doc, "Persistența și documentele")
    add_bullets(doc, [
        "MariaDB: stocarea datelor operaționale și CRM;",
        "MariaDB JDBC și HikariCP: acces direct și reutilizarea conexiunilor;",
        "repository-uri și DAO-uri scrise manual: păstrarea SQL-ului explicit, fără ORM;",
        "Apache POI: import și export Excel;",
        "OpenPDF: facturi și exporturi PDF.",
    ])

    heading(doc, "Serviciul AI")
    add_body(doc, "Anthropic Claude este integrat prin SDK-ul Java și furnizează asistentul, recomandările și traducerea. Clientul este inițializat numai dacă există o cheie API; în lipsa acesteia, nivelul AI este dezactivat controlat.")

    chapter(doc, 8, "Implementarea și modelul de date")
    heading(doc, "Organizarea codului")
    add_body(doc, "Componenta server conține aproximativ 11.500 de linii de cod Java, organizate pe responsabilități: builderi, comenzi, DAO-uri, fabrici, modele, observatori, șabloane, repository-uri, servicii, strategii, validare și nivel web. Clientul cuprinde aproximativ 50 de module TypeScript și React.")

    heading(doc, "Nucleul bazat pe evenimente")
    add_body(doc, "Comportamentul «o acțiune, mai multe reacții» este implementat prin șablonul Observer și EventBus. După o operație importantă, serviciul publică un eveniment, iar observatorii înregistrați execută independent reacțiile aferente.")
    add_bullets(doc, [
        "înregistrarea în jurnalul de audit;",
        "trimiterea mesajelor de bun venit și confirmare;",
        "generarea facturilor;",
        "actualizarea scorurilor clienților potențiali.",
    ])
    add_body(doc, "Observatorul de facturare ascultă evenimentul de rezervare, verifică ședința și solicită generarea facturii. Erorile sunt înregistrate local, astfel încât rezervarea să nu fie anulată.")
    add_key_point(doc, "Idempotent", "Dacă există deja o factură pentru ședință, sistemul nu creează alta, chiar dacă evenimentul este repetat.")

    heading(doc, "Evenimentele domeniului")
    add_bullets(doc, [
        "crearea unui contact;",
        "crearea unei înscrieri;",
        "rezervarea unei ședințe;",
        "finalizarea unei activități;",
        "schimbarea stării unui lead;",
        "schimbarea etapei unei oportunități.",
    ])

    heading(doc, "Modelul de date")
    add_body(doc, "Modelul integrează tabelele CRM și tabelele funcționalităților web. Contact poate reprezenta o persoană INDIVIDUAL sau o companie CORPORATE, iar companiile sunt asociate angajaților lor.")
    add_bullets(doc, [
        "Course oferă una sau mai multe CourseSession;",
        "Enrollment leagă un contact de sesiunea unui curs;",
        "Opportunity și Activity descriu procesul comercial;",
        "trainerii susțin ședințe individuale, care generează facturi.",
    ])

    heading(doc, "Stocarea recenziilor")
    add_body(doc, "Recenziile sunt stocate direct în Enrollment prin câmpurile rating și feedback. Această structură permite o singură recenzie pentru fiecare pereche client-curs; retrimiterea înlocuiește valorile anterioare. Valorile 1-5 sunt evaluări valide, iar 0 înseamnă că acel curs nu a fost încă evaluat.")

    heading(doc, "Integrarea Claude")
    add_body(doc, "ClaudeClient este un strat subțire peste SDK-ul oficial Anthropic. O singură cheie activează asistentul conversațional, chestionarul de recomandare, recomandările pentru companii și traducerea în timp real.")
    add_body(doc, "Dacă cheia lipsește, clientul nu este inițializat, isEnabled() returnează false, iar endpointurile raportează că funcția este oprită fără a produce blocarea aplicației.")

    heading(doc, "Statistici și documente")
    add_body(doc, "Statisticile publice sunt transmise prin SSE. Broadcasterul recalculează valorile cât timp există cel puțin un client conectat și transmite numai modificările, evitând interogările periodice.")
    add_body(doc, "Apache POI este utilizat pentru fișiere Excel, iar OpenPDF pentru facturi și exporturi PDF.")

    heading(doc, "Clarificare: ordinea contact-înscriere")
    add_numbered(doc, [
        "La înregistrare, utilizatorul este creat ca contact în CRM.",
        "La cumpărare, sistemul caută contactul după e-mail.",
        "Dacă lipsește, îl creează înainte de continuarea fluxului.",
        "Sistemul identifică sau creează sesiunea cursului.",
        "Abia apoi creează Enrollment și publică evenimentul.",
    ])
    add_key_point(doc, "Concluzie", "Contactul nu este creat după înscriere; el trebuie să existe înainte pentru ca Enrollment să poată fi asociat corect.")

    chapter(doc, 9, "Testarea aplicației")
    heading(doc, "Strategia de testare")
    add_body(doc, "Corectitudinea TrainingIT depinde de propagarea corectă a acțiunilor prin întregul sistem. Testarea s-a concentrat asupra validării datelor, robusteții nucleului bazat pe evenimente și verificării fluxurilor complete.")

    heading(doc, "Validarea pe două niveluri")
    add_bullets(doc, [
        "clientul verifică existența câmpurilor, confirmarea parolei și consimțământul GDPR;",
        "serverul folosește Spring Validation pentru cererile primite;",
        "datele contactelor trec printr-un lanț Chain of Responsibility care verifică e-mailul, telefonul, GDPR-ul și tipul contactului.",
    ])

    heading(doc, "Comportamentul în caz de eroare")
    add_bullets(doc, [
        "un eveniment repetat nu generează facturi duplicate;",
        "eroarea unui observator nu anulează rezervarea;",
        "lipsa cheii API ascunde funcțiile AI fără a afecta sistemul;",
        "ștergerea unei sesizări necesită confirmare explicită.",
    ])

    heading(doc, "Testarea funcțională manuală")
    add_body(doc, "Fluxurile principale au fost testate pe o instanță funcțională conectată la o bază de date MariaDB reală.")
    add_bullets(doc, [
        "înregistrarea și crearea leadului;",
        "direcționarea rolurilor USER și ADMIN;",
        "cumpărarea și apariția imediată în Purchases;",
        "recenzia și actualizarea evaluării medii;",
        "rezervarea, factura și reducerea angajatului;",
        "trimiterea sesizării către Issues;",
        "exportarea contactelor în CSV, Excel și PDF.",
    ])
    add_body(doc, "Proiectul include spring-boot-starter-test ca bază pentru dezvoltarea viitoare a unei suite automate de teste unitare și de integrare.")

    chapter(doc, 10, "Scenarii de utilizare")
    heading(doc, "Alexandru - cursant individual")
    add_body(doc, "Alexandru dorește să treacă din vânzări în IT. Asistentul îi recomandă două cursuri pentru începători. El se înregistrează, devine lead în CRM, cumpără un curs și programează o ședință individuală, pentru care sistemul emite automat factura.")
    add_key_point(doc, "Rezultat", "Un vizitator nehotărât devine rapid client plătitor, fără introducerea manuală a datelor de către companie.")

    heading(doc, "Cristina - client venit prin recomandare")
    add_body(doc, "Cristina este dezvoltator junior și ajunge prin recomandarea unui prieten, ceea ce îi oferă un scor ridicat. Ea solicită cursuri avansate, află despre ședințele individuale și cumpără rapid.")
    add_key_point(doc, "Rezultat", "Recomandările cursanților mulțumiți au un cost redus de achiziție și sunt vizibile în analiza surselor leadurilor.")

    heading(doc, "BankTech Solutions - client corporativ")
    add_body(doc, "Administratorul importă angajații companiei dintr-un fișier Excel și solicită recomandări AI pentru întreaga echipă. CRM-ul creează o oportunitate B2B, angajații beneficiază de reducerea automată de 60%, iar ședințele individuale produc facturi fără intervenție manuală.")
    add_key_point(doc, "Rezultat", "O singură companie aduce simultan numeroși cursanți, iar CRM-ul urmărește echipa, reducerile și facturile.")

    heading(doc, "Fluxul comun de cumpărare și recenzie")
    add_numbered(doc, [
        "Identificarea sau crearea contactului.",
        "Identificarea sau crearea sesiunii cursului.",
        "Salvarea Enrollment.",
        "Publicarea evenimentului de înscriere.",
        "Recalcularea și transmiterea statisticilor.",
        "Validarea achiziției înaintea recenziei.",
        "Salvarea evaluării și comentariului în Enrollment.",
    ])

    chapter(doc, 11, "Concluzii și dezvoltări viitoare")
    heading(doc, "Rezultatul lucrării")
    add_body(doc, "TrainingIT reunește într-un singur produs o platformă de comercializare a cursurilor și un sistem CRM complet. Cele două portaluri folosesc același backend și același model de date, astfel încât fiecare acțiune este înregistrată o singură dată și declanșează automat procesele corespunzătoare.")
    add_body(doc, "La înregistrare, vizitatorul devine un lead urmărit în CRM, iar la prima achiziție devine client. Acțiunile ulterioare sunt reflectate automat în portalul administrativ.")

    heading(doc, "Îndeplinirea obiectivelor")
    add_bullets(doc, [
        "portal intuitiv și ușor de utilizat pentru clienți;",
        "portal administrativ în șapte secțiuni;",
        "nivel de domeniu bazat pe șabloane clasice;",
        "nucleu bazat pe evenimente și reacții automate;",
        "persistența datelor în MariaDB;",
        "funcții AI și traducere disponibile opțional.",
    ])
    add_key_point(doc, "Caracter distinctiv", "Aplicația combină o interfață intuitivă cu un nucleu CRM bine structurat și un nivel opțional de inteligență artificială.")

    heading(doc, "Direcții viitoare")
    add_bullets(doc, [
        "testare automată: teste unitare, de integrare și integrare continuă;",
        "autentificare mai sigură: tokenuri, hashing al parolelor, verificarea e-mailului și resetare prin linkuri temporare;",
        "analize mai complexe: tendințe istorice, cohorte și dashboarduri exportabile;",
        "funcționalități AI extinse: context mai bogat, memorie conversațională și rezumarea sesizărilor;",
        "implementare și scalare: containerizare, configurație externă și pregătire pentru cloud.",
    ])

    chapter(doc, "A", "Glosar de termeni")
    add_term_table(doc, [
        ("CRM", "Sistem pentru gestionarea contactelor, leadurilor, vânzărilor, facturilor și interacțiunilor cu clienții."),
        ("Platformă agregatoare", "Platformă care reunește ofertele mai multor furnizori și intermediază relația cu utilizatorii, precum Udemy sau Coursera."),
        ("Jurnal de audit", "Evidență cronologică a acțiunilor importante: cine a făcut acțiunea, când și asupra cărei înregistrări."),
        ("Logică de afaceri", "Regulile care stabilesc cum funcționează aplicația: reduceri, disponibilitate, facturare, scoruri și schimbări de stare."),
        ("Full-stack", "Sistem care include interfața, backendul, baza de date și comunicarea dintre acestea."),
        ("REST API", "Interfață prin care front-endul trimite cereri către backend și primește date sau rezultate."),
        ("SSE", "Server-Sent Events; mecanism prin care serverul transmite actualizări în timp real către browser."),
        ("DTO", "Obiect folosit pentru transportul datelor între client, controllere și servicii."),
        ("DAO", "Componentă care execută operațiile directe de acces la baza de date."),
        ("Repository", "Componentă care oferă operații de persistență într-o formă apropiată de domeniul aplicației."),
        ("Observer", "Șablon prin care o componentă reacționează la evenimente fără a fi legată direct de operația care le-a produs."),
        ("EventBus", "Mecanism central care publică evenimente și notifică observatorii înregistrați."),
        ("Idempotent", "O operație repetată produce același rezultat și nu creează efecte duplicate."),
        ("Spring Validation", "Mecanism Spring pentru verificarea declarativă a datelor primite de backend."),
        ("Chain of Responsibility", "Șablon care transmite datele printr-un lanț de verificări, fiecare componentă având o responsabilitate distinctă."),
        ("Robustețe", "Capacitatea sistemului de a rămâne stabil și previzibil atunci când apar erori sau evenimente neprevăzute."),
        ("Degradare controlată", "Dezactivarea unei funcții opționale fără blocarea sau afectarea celorlalte funcționalități."),
        ("Interfață intuitivă", "Interfață clară și ușor de folosit, inclusiv de persoane fără pregătire tehnică; termenul nu afirmă automat conformitatea WCAG."),
    ])

    # Metadata and save.
    props = doc.core_properties
    props.title = "Conspect TrainingIT"
    props.subject = "Sinteza în limba română a lucrării TrainingIT"
    props.author = ""
    props.keywords = "TrainingIT, CRM, cursuri online, conspect, licență"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
